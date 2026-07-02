"""
Generate Supplementary Document: Mathematical Framework
=========================================================
Comprehensive documentation of all mathematical expressions used in the
semi-mechanistic PBPK/PD model for mycophenolic acid (MPA).
Professional OMML equations rendered in DOCX format.
"""

import os
from lxml import etree
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

try:
    from generate_updated_manuscript_docx import (
        build_references,
        write_ris,
        _endnote_temp_marker,
    )
except ImportError:
    from simulations.generate_updated_manuscript_docx import (
        build_references,
        write_ris,
        _endnote_temp_marker,
    )

project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG_DIR = os.path.join(project_root, "outputs", "manuscript_updated")
OUTPUT_PATH = os.path.join(FIG_DIR, "Supplement_Mathematical_Framework_endnote_ready.docx")
RIS_PATH = os.path.join(FIG_DIR, "MPA_QSP_References.ris")

MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
WORD_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


REFERENCES = build_references()


def cite(*ref_nums):
    return _endnote_temp_marker(ref_nums, REFERENCES)


VISIBLE_NOTATION_REPLACEMENTS = {
    'AUC₀₋₁₂ₕ': 'AUC_0-12h',
    'AUC₀₋τ': 'AUC_0-tau',
    'AUCᶠᵣᵉᵉ': 'AUC_free',
    'AUCₜₒₜₐₗ': 'AUC_total',
    'AUC₅₀': 'AUC50',
    'IMPDH₅₀': 'IMPDH50',
    'CLᵢₙₜ': 'CL_int',
    'CLₕ': 'CL_h',
    'CLᵣ': 'CL_r',
    'Qₕ': 'Q_h',
    'Vₘₚₐᵍ': 'V_MPAG',
    'Vᶜ': 'V_c',
    'Vₚ': 'V_p',
    'Cₘₚₐᵍ': 'C_MPAG',
    'Cₘₚₐ': 'C_MPA',
    'Cᶠᵣᵉᵉ': 'C_free',
    'Cₘₐₓ': 'C_max',
    'tₘₐₓ': 't_max',
    'MWₘₚₐᵍ': 'MW_MPAG',
    'MWₘₚₐ': 'MW_MPA',
    'MWₘₘᶠ': 'MW_MMF',
    'Albᵣᵉᶠ': 'Alb_ref',
    'GFRᵣᵉᶠ': 'GFR_ref',
    'LWᵣₑƒ': 'LW_ref',
    'LWᵢ': 'LW_i',
    'fₒₐₜₚ₁ₑ': 'f_OATP1B',
    'fᵤ': 'f_u',
    'f̅ᵤ': 'f_u_avg',
    'fₘ': 'f_m',
    'Fₒᵣₐₗ': 'F_oral',
    'Fᵍ₁': 'F_gw',
    'Fᶜₒₙᵥ': 'F_conv',
    'kₐ': 'k_a',
    'kₑᵤₜ': 'k_gut',
    'kₑᵢₗᵉ': 'k_bile',
    'kₚᵣₒₗᵢᶠ': 'k_prolif',
    'kᵈᵉₐₜₕ': 'k_death',
    'tₗₐᵍ': 't_lag',
    'Aₑᵤₜ': 'A_gut',
    'Aₑᵢₗᵉ': 'A_bile',
    'Aₐₑₛ': 'A_abs',
    'Aₚ': 'A_p',
    'Iₐₑᶜᶜ₂': 'I_ABCC2',
    'Iₒₐₜₚ': 'I_OATP',
    'Iₐᵥᵍ': 'I_avg',
    'IC₅₀': 'IC50',
    'EC₅₀': 'EC50',
    'Eₘₐₓ': 'E_max',
    'Kᵢ': 'K_i',
    'L₀': 'L_0',
    'Lₛₛ': 'L_ss',
    'Lₜₕᵣᵉₛₕ': 'L_thresh',
    'Pₑₐₛᵉ': 'P_base',
    'Pₘₐₓ': 'P_max',
    'Pₐᵈᵥᵉᵣₛᵉ': 'P_adverse',
    'Pᵣᵉⱼ': 'P_rej',
    'Pᵍᴵ': 'P_GI',
    'Pₗᵉᵤₖₒ': 'P_leuko',
    'Pᵢₙᶠ': 'P_inf',
    'γ₁': 'gamma_1',
    'γ₂': 'gamma_2',
    'CL_int,ᵣₑƒ': 'CL_int_ref',
    'CL_int,ᵢ': 'CL_int_i',
    '(BW/70)⁰·⁷⁵': '(BW/70)^0.75',
    'BW⁰·⁸⁶': 'BW^0.86',
    '⁰⋈⁸⁶': '^0.86',
    '⁰·⁷⁵': '^0.75',
    'ᵣₑƒ': '_ref',
    'ᵢ': '_i',
    'm²': 'm^2',
    '10⁹': '10^9',
    '10⁻⁸': '10^-8',
    '10⁻¹⁰': '10^-10',
}


def normalize_visible_notation(doc):
    """Convert hard-to-read Unicode notation in prose/tables to ASCII labels.

    OMML equation objects are left untouched; only ordinary text runs are changed.
    """
    replacements = sorted(
        VISIBLE_NOTATION_REPLACEMENTS.items(),
        key=lambda item: len(item[0]),
        reverse=True,
    )

    def normalize_text(text):
        for old, new in replacements:
            text = text.replace(old, new)
        return text

    def normalize_paragraph(paragraph):
        for run in paragraph.runs:
            new_text = normalize_text(run.text)
            if new_text != run.text:
                run.text = new_text

    for paragraph in doc.paragraphs:
        normalize_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    normalize_paragraph(paragraph)


# ================================================================
# OMML HELPERS
# ================================================================

def math_sub(parent, tag):
    return etree.SubElement(parent, f'{{{MATH_NS}}}{tag}')

def word_sub(parent, tag):
    return etree.SubElement(parent, f'{{{WORD_NS}}}{tag}')

def math_run(parent, text, italic=True, bold=False, font_name='Cambria Math'):
    r = math_sub(parent, 'r')
    if not italic or bold:
        rPr = math_sub(r, 'rPr')
        if not italic:
            sty = math_sub(rPr, 'sty')
            sty.set(f'{{{MATH_NS}}}val', 'p')
    w_rPr = word_sub(r, 'rPr')
    rFonts = word_sub(w_rPr, 'rFonts')
    rFonts.set(f'{{{WORD_NS}}}ascii', font_name)
    rFonts.set(f'{{{WORD_NS}}}hAnsi', font_name)
    if bold:
        word_sub(w_rPr, 'b')
    t = math_sub(r, 't')
    t.text = text
    return r

def math_fraction(parent, bar_type='bar'):
    f = math_sub(parent, 'f')
    fPr = math_sub(f, 'fPr')
    typ = math_sub(fPr, 'type')
    typ.set(f'{{{MATH_NS}}}val', bar_type)
    num = math_sub(f, 'num')
    den = math_sub(f, 'den')
    return f, num, den

def math_subscript(parent, base_text, sub_text, base_italic=True, sub_italic=True):
    sSub = math_sub(parent, 'sSub')
    math_sub(sSub, 'sSubPr')
    e = math_sub(sSub, 'e')
    math_run(e, base_text, italic=base_italic)
    sub = math_sub(sSub, 'sub')
    math_run(sub, sub_text, italic=sub_italic)
    return sSub

def math_superscript(parent, base_text, sup_text, base_italic=True):
    sSup = math_sub(parent, 'sSup')
    math_sub(sSup, 'sSupPr')
    e = math_sub(sSup, 'e')
    math_run(e, base_text, italic=base_italic)
    sup = math_sub(sSup, 'sup')
    math_run(sup, sup_text)
    return sSup

def math_subsup(parent, base_text, sub_text, sup_text, base_italic=True):
    sSubSup = math_sub(parent, 'sSubSup')
    math_sub(sSubSup, 'sSubSupPr')
    e = math_sub(sSubSup, 'e')
    math_run(e, base_text, italic=base_italic)
    sub_elem = math_sub(sSubSup, 'sub')
    math_run(sub_elem, sub_text)
    sup_elem = math_sub(sSubSup, 'sup')
    math_run(sup_elem, sup_text)
    return sSubSup

def math_delimiters(parent, open_char='(', close_char=')'):
    d = math_sub(parent, 'd')
    dPr = math_sub(d, 'dPr')
    begChr = math_sub(dPr, 'begChr')
    begChr.set(f'{{{MATH_NS}}}val', open_char)
    endChr = math_sub(dPr, 'endChr')
    endChr.set(f'{{{MATH_NS}}}val', close_char)
    e = math_sub(d, 'e')
    return e

def math_nary(parent, char='\u2211', sub_text=None, sup_text=None):
    """Create a big operator (sum, product, integral)."""
    nary = math_sub(parent, 'nary')
    naryPr = math_sub(nary, 'naryPr')
    chr_elem = math_sub(naryPr, 'chr')
    chr_elem.set(f'{{{MATH_NS}}}val', char)
    limLoc = math_sub(naryPr, 'limLoc')
    limLoc.set(f'{{{MATH_NS}}}val', 'subSup')
    sub_e = math_sub(nary, 'sub')
    if sub_text:
        math_run(sub_e, sub_text, italic=False)
    sup_e = math_sub(nary, 'sup')
    if sup_text:
        math_run(sup_e, sup_text, italic=False)
    body_e = math_sub(nary, 'e')
    return body_e

def create_display_equation(doc, builder_func, number=None):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    oMathPara = etree.SubElement(para._element, f'{{{MATH_NS}}}oMathPara')
    oMath = math_sub(oMathPara, 'oMath')
    builder_func(oMath)
    if number:
        run = para.add_run(f'\t({number})')
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
    return para

def add_inline_math(paragraph, builder_func):
    oMath = etree.SubElement(paragraph._element, f'{{{MATH_NS}}}oMath')
    builder_func(oMath)
    return oMath


# ================================================================
# EQUATION BUILDERS - S1. HEPATIC CLEARANCE MODEL
# ================================================================

def eq_well_stirred(oMath):
    math_subscript(oMath, 'CL', 'h')
    math_run(oMath, ' = ', italic=False)
    f, num, den = math_fraction(oMath)
    math_subscript(num, 'Q', 'h')
    math_run(num, ' \u00b7 ')
    math_subscript(num, 'f', 'u')
    math_run(num, ' \u00b7 ')
    math_subscript(num, 'CL', 'int')
    math_subscript(den, 'Q', 'h')
    math_run(den, ' + ')
    math_subscript(den, 'f', 'u')
    math_run(den, ' \u00b7 ')
    math_subscript(den, 'CL', 'int')

def eq_extraction_ratio(oMath):
    math_run(oMath, 'E')
    math_run(oMath, ' = ', italic=False)
    f, num, den = math_fraction(oMath)
    math_subscript(num, 'f', 'u')
    math_run(num, ' \u00b7 ')
    math_subscript(num, 'CL', 'int')
    math_subscript(den, 'Q', 'h')
    math_run(den, ' + ')
    math_subscript(den, 'f', 'u')
    math_run(den, ' \u00b7 ')
    math_subscript(den, 'CL', 'int')

def eq_low_e_approx(oMath):
    math_subscript(oMath, 'CL', 'h')
    math_run(oMath, ' \u2248 ', italic=False)
    math_subscript(oMath, 'f', 'u')
    math_run(oMath, ' \u00b7 ')
    math_subscript(oMath, 'CL', 'int')
    math_run(oMath, '          (when ', italic=False)
    math_run(oMath, 'E')
    math_run(oMath, ' \u226a 1)', italic=False)

def eq_total_clint(oMath):
    math_subscript(oMath, 'CL', 'int')
    math_run(oMath, ' = ', italic=False)
    sSub1 = math_sub(oMath, 'sSub')
    math_sub(sSub1, 'sSubPr')
    e1 = math_sub(sSub1, 'e')
    math_run(e1, 'CL')
    sub1 = math_sub(sSub1, 'sub')
    math_run(sub1, 'int,UGT1A9', italic=False)
    math_run(oMath, ' + ', italic=False)
    sSub2 = math_sub(oMath, 'sSub')
    math_sub(sSub2, 'sSubPr')
    e2 = math_sub(sSub2, 'e')
    math_run(e2, 'CL')
    sub2 = math_sub(sSub2, 'sub')
    math_run(sub2, 'int,UGT2B7', italic=False)


# ================================================================
# EQUATION BUILDERS - S2. PROTEIN BINDING
# ================================================================

def eq_fu_albumin(oMath):
    math_subscript(oMath, 'f', 'u,base', sub_italic=False)
    math_run(oMath, ' = ', italic=False)
    sSub = math_sub(oMath, 'sSub')
    math_sub(sSub, 'sSubPr')
    e = math_sub(sSub, 'e')
    math_run(e, 'f')
    sub = math_sub(sSub, 'sub')
    math_run(sub, 'u,ref', italic=False)
    math_run(oMath, ' \u00b7 ', italic=False)
    sSup = math_sub(oMath, 'sSup')
    math_sub(sSup, 'sSupPr')
    e2 = math_sub(sSup, 'e')
    inner = math_delimiters(e2)
    f, num, den = math_fraction(inner)
    math_subscript(num, 'Alb', 'ref')
    math_run(den, 'Alb')
    sup = math_sub(sSup, 'sup')
    math_run(sup, '1.2', italic=False)

def eq_mpag_displacement(oMath):
    math_subscript(oMath, 'f', 'disp')
    math_run(oMath, ' = 1 + ', italic=False)
    f, num, den = math_fraction(oMath)
    math_run(num, '0.50 \u00b7 ', italic=False)
    math_subscript(num, 'C', 'MPAG')
    math_run(den, '30 + ', italic=False)
    math_subscript(den, 'C', 'MPAG')

def eq_fu_combined(oMath):
    math_subscript(oMath, 'f', 'u')
    math_run(oMath, ' = ', italic=False)
    math_subscript(oMath, 'f', 'u,base', sub_italic=False)
    math_run(oMath, ' \u00b7 ')
    math_subscript(oMath, 'f', 'disp')
    math_run(oMath, ',     ', italic=False)
    math_subscript(oMath, 'f', 'u')
    math_run(oMath, ' \u2208 [0.005, 0.30]', italic=False)


# ================================================================
# EQUATION BUILDERS - S3. ALLOMETRIC SCALING
# ================================================================

def eq_allometric_volume(oMath):
    math_subscript(oMath, 'V', 'i')
    math_run(oMath, ' = ', italic=False)
    math_subscript(oMath, 'V', 'ref')
    math_run(oMath, ' \u00b7 ', italic=False)
    sSup = math_sub(oMath, 'sSup')
    math_sub(sSup, 'sSupPr')
    e = math_sub(sSup, 'e')
    inner = math_delimiters(e)
    f, num, den = math_fraction(inner)
    math_subscript(num, 'BW', 'i')
    math_subscript(den, 'BW', 'ref')
    sup = math_sub(sSup, 'sup')
    math_run(sup, '1.0', italic=False)

def eq_allometric_clearance(oMath):
    math_subscript(oMath, 'CL', 'i')
    math_run(oMath, ' = ', italic=False)
    math_subscript(oMath, 'CL', 'ref')
    math_run(oMath, ' \u00b7 ', italic=False)
    sSup = math_sub(oMath, 'sSup')
    math_sub(sSup, 'sSupPr')
    e = math_sub(sSup, 'e')
    inner = math_delimiters(e)
    f, num, den = math_fraction(inner)
    math_subscript(num, 'BW', 'i')
    math_subscript(den, 'BW', 'ref')
    sup = math_sub(sSup, 'sup')
    math_run(sup, '0.75', italic=False)

def eq_liver_weight(oMath):
    math_subscript(oMath, 'LW', 'i')
    math_run(oMath, ' = 1.5 \u00b7 ', italic=False)
    sSup = math_sub(oMath, 'sSup')
    math_sub(sSup, 'sSupPr')
    e = math_sub(sSup, 'e')
    inner = math_delimiters(e)
    f, num, den = math_fraction(inner)
    math_subscript(num, 'BW', 'i')
    math_run(den, '70', italic=False)
    sup = math_sub(sSup, 'sup')
    math_run(sup, '0.86', italic=False)
    math_run(oMath, '     kg', italic=False)

def eq_hepatic_blood_flow(oMath):
    math_subscript(oMath, 'Q', 'h,i')
    math_run(oMath, ' = 90 \u00b7 ', italic=False)
    sSup = math_sub(oMath, 'sSup')
    math_sub(sSup, 'sSupPr')
    e = math_sub(sSup, 'e')
    inner = math_delimiters(e)
    f, num, den = math_fraction(inner)
    math_subscript(num, 'BW', 'i')
    math_run(den, '70', italic=False)
    sup = math_sub(sSup, 'sup')
    math_run(sup, '0.75', italic=False)
    math_run(oMath, '     L/h', italic=False)

def eq_bsa(oMath):
    math_run(oMath, 'BSA', italic=False)
    math_run(oMath, ' = 0.007184 \u00b7 ', italic=False)
    math_superscript(oMath, 'H', '0.725')
    math_run(oMath, ' \u00b7 ')
    math_superscript(oMath, 'BW', '0.425')

def eq_renal_clearance(oMath):
    sSub = math_sub(oMath, 'sSub')
    math_sub(sSub, 'sSubPr')
    e = math_sub(sSub, 'e')
    math_run(e, 'CL')
    sub = math_sub(sSub, 'sub')
    math_run(sub, 'r,MPAG', italic=False)
    math_run(oMath, ' = ', italic=False)
    sSub2 = math_sub(oMath, 'sSub')
    math_sub(sSub2, 'sSubPr')
    e2 = math_sub(sSub2, 'e')
    math_run(e2, 'CL')
    sub2 = math_sub(sSub2, 'sub')
    math_run(sub2, 'r,ref', italic=False)
    math_run(oMath, ' \u00b7 ', italic=False)
    f, num, den = math_fraction(oMath)
    math_run(num, 'GFR')
    math_subscript(den, 'GFR', 'ref')

def eq_gut_wall(oMath):
    math_subscript(oMath, 'F', 'gw')
    math_run(oMath, ' = ', italic=False)
    f, num, den = math_fraction(oMath)
    math_subscript(num, 'F', 'gw,ref', sub_italic=False)
    math_run(den, 'max', italic=False)
    inner = math_delimiters(den)
    math_subscript(inner, 'A', 'UGT1A9', sub_italic=False)
    math_run(inner, ', 0.5', italic=False)


# ================================================================
# EQUATION BUILDERS - S4. ODE SYSTEM
# ================================================================

def eq_ode_gut(oMath):
    f_lhs, num_lhs, den_lhs = math_fraction(oMath)
    math_run(num_lhs, 'd')
    math_subscript(num_lhs, 'A', 'gut')
    math_run(den_lhs, 'dt')
    math_run(oMath, ' = \u2212', italic=False)
    math_subscript(oMath, 'k', 'a')
    math_run(oMath, ' \u00b7 ')
    math_subscript(oMath, 'A', 'gut')

def eq_ode_central(oMath):
    f_lhs, num_lhs, den_lhs = math_fraction(oMath)
    math_run(num_lhs, 'd')
    math_subscript(num_lhs, 'C', 'MPA')
    math_run(den_lhs, 'dt')
    math_run(oMath, ' = ', italic=False)
    f_rhs, num_rhs, den_rhs = math_fraction(oMath)
    math_subscript(num_rhs, 'R', 'abs')
    math_run(num_rhs, ' + ')
    math_subscript(num_rhs, 'R', 'dist,in')
    math_run(num_rhs, ' + ')
    math_subscript(num_rhs, 'R', 'EHC')
    math_run(num_rhs, ' \u2212 ')
    math_subscript(num_rhs, 'R', 'dist,out')
    math_run(num_rhs, ' \u2212 ')
    math_subscript(num_rhs, 'R', 'met')
    math_subscript(den_rhs, 'V', 'c')

def eq_ode_periph(oMath):
    f_lhs, num_lhs, den_lhs = math_fraction(oMath)
    math_run(num_lhs, 'd')
    math_subscript(num_lhs, 'A', 'p')
    math_run(den_lhs, 'dt')
    math_run(oMath, ' = ', italic=False)
    math_run(oMath, 'Q')
    math_run(oMath, ' \u00b7 ')
    math_subscript(oMath, 'C', 'MPA')
    math_run(oMath, ' \u2212 ')
    math_run(oMath, 'Q')
    math_run(oMath, ' \u00b7 ')
    f, num, den = math_fraction(oMath)
    math_subscript(num, 'A', 'p')
    math_subscript(den, 'V', 'p')

def eq_ode_mpag(oMath):
    f_lhs, num_lhs, den_lhs = math_fraction(oMath)
    math_run(num_lhs, 'd')
    math_subscript(num_lhs, 'C', 'MPAG')
    math_run(den_lhs, 'dt')
    math_run(oMath, ' = ', italic=False)
    f_rhs, num_rhs, den_rhs = math_fraction(oMath)
    math_subscript(num_rhs, 'CL', 'h')
    math_run(num_rhs, ' \u00b7 ')
    math_subscript(num_rhs, 'C', 'MPA')
    math_run(num_rhs, ' \u2212 ')
    sSub1 = math_sub(num_rhs, 'sSub')
    math_sub(sSub1, 'sSubPr')
    e1 = math_sub(sSub1, 'e')
    math_run(e1, 'CL')
    sub1 = math_sub(sSub1, 'sub')
    math_run(sub1, 'r', italic=False)
    math_run(num_rhs, ' \u00b7 ')
    math_subscript(num_rhs, 'C', 'MPAG')
    math_run(num_rhs, ' \u2212 ')
    math_subscript(num_rhs, 'R', 'bile')
    math_subscript(den_rhs, 'V', 'MPAG')

def eq_ode_bile(oMath):
    f_lhs, num_lhs, den_lhs = math_fraction(oMath)
    math_run(num_lhs, 'd')
    math_subscript(num_lhs, 'A', 'bile')
    math_run(den_lhs, 'dt')
    math_run(oMath, ' = ', italic=False)
    math_subscript(oMath, 'R', 'bile')
    math_run(oMath, ' \u2212 ')
    math_subscript(oMath, 'k', 'gut')
    math_run(oMath, ' \u00b7 ')
    math_subscript(oMath, 'A', 'bile')


# ================================================================
# EQUATION BUILDERS - S5. RATE EXPRESSIONS
# ================================================================

def eq_rate_absorption(oMath):
    math_subscript(oMath, 'R', 'abs')
    math_run(oMath, ' = ', italic=False)
    math_subscript(oMath, 'k', 'a')
    math_run(oMath, ' \u00b7 ')
    math_subscript(oMath, 'A', 'gut')
    math_run(oMath, ' \u00b7 ')
    math_subscript(oMath, 'F', 'oral')
    math_run(oMath, ' \u00b7 ')
    math_subscript(oMath, 'F', 'gw')

def eq_rate_distribution(oMath):
    math_subscript(oMath, 'R', 'dist,out')
    math_run(oMath, ' = ', italic=False)
    math_run(oMath, 'Q')
    math_run(oMath, ' \u00b7 ')
    math_subscript(oMath, 'C', 'MPA')
    math_run(oMath, ',      ', italic=False)
    math_subscript(oMath, 'R', 'dist,in')
    math_run(oMath, ' = ', italic=False)
    math_run(oMath, 'Q')
    math_run(oMath, ' \u00b7 ')
    f, num, den = math_fraction(oMath)
    math_subscript(num, 'A', 'p')
    math_subscript(den, 'V', 'p')

def eq_rate_metabolism(oMath):
    math_subscript(oMath, 'R', 'met')
    math_run(oMath, ' = ', italic=False)
    math_subscript(oMath, 'CL', 'h')
    math_run(oMath, ' \u00b7 ')
    math_subscript(oMath, 'C', 'MPA')

def eq_rate_ehc(oMath):
    math_subscript(oMath, 'R', 'EHC')
    math_run(oMath, ' = ', italic=False)
    math_subscript(oMath, 'k', 'gut')
    math_run(oMath, ' \u00b7 ')
    math_subscript(oMath, 'A', 'bile')
    math_run(oMath, ' \u00b7 ')
    math_subscript(oMath, 'f', 'EHC')

def eq_rate_biliary(oMath):
    math_subscript(oMath, 'R', 'bile')
    math_run(oMath, ' = ', italic=False)
    math_subscript(oMath, 'k', 'bile')
    math_run(oMath, ' \u00b7 ')
    math_subscript(oMath, 'C', 'MPAG')
    math_run(oMath, ' \u00b7 ')
    math_subscript(oMath, 'V', 'MPAG')
    math_run(oMath, ' \u00b7 ')
    math_subscript(oMath, 'f', 'OATP1B')


# ================================================================
# EQUATION BUILDERS - S6. DRUG INTERACTIONS
# ================================================================

def eq_csa_ehc(oMath):
    math_subscript(oMath, 'f', 'EHC,CsA', sub_italic=False)
    math_run(oMath, ' = ', italic=False)
    math_subscript(oMath, 'f', 'EHC')
    math_run(oMath, ' \u00b7 ', italic=False)
    inner = math_delimiters(oMath)
    math_run(inner, '1 \u2212 ', italic=False)
    math_subscript(inner, 'I', 'ABCC2')

def eq_csa_oatp(oMath):
    math_subscript(oMath, 'f', 'OATP1B,CsA', sub_italic=False)
    math_run(oMath, ' = ', italic=False)
    math_subscript(oMath, 'f', 'OATP1B')
    math_run(oMath, ' \u00b7 ', italic=False)
    inner = math_delimiters(oMath)
    math_run(inner, '1 \u2212 ', italic=False)
    math_subscript(inner, 'I', 'OATP')


# ================================================================
# EQUATION BUILDERS - S7. DOSE & EXPOSURE
# ================================================================

def eq_mpa_dose(oMath):
    math_subscript(oMath, 'D', 'MPA')
    math_run(oMath, ' = ', italic=False)
    math_subscript(oMath, 'D', 'MMF')
    math_run(oMath, ' \u00b7 ', italic=False)
    f, num, den = math_fraction(oMath)
    math_subscript(num, 'MW', 'MPA')
    math_subscript(den, 'MW', 'MMF')
    math_run(oMath, ' \u00b7 ')
    math_subscript(oMath, 'F', 'conv')

def eq_auc_trapezoidal(oMath):
    math_subscript(oMath, 'AUC', '0\u2013\u03c4')
    math_run(oMath, ' = ', italic=False)
    body = math_nary(oMath, '\u222b', '0', '\u03c4')
    math_run(body, 'C')
    inner = math_delimiters(body)
    math_run(inner, 't')
    math_run(body, ' dt')

def eq_auc_free(oMath):
    math_subscript(oMath, 'AUC', 'free')
    math_run(oMath, ' = ', italic=False)
    body = math_nary(oMath, '\u222b', '0', '\u03c4')
    math_subscript(body, 'f', 'u')
    inner = math_delimiters(body)
    math_run(inner, 't')
    math_run(body, ' \u00b7 C')
    inner2 = math_delimiters(body)
    math_run(inner2, 't')
    math_run(body, ' dt')

def eq_auc_target(oMath):
    math_run(oMath, '30', italic=False)
    math_run(oMath, ' \u2264 ', italic=False)
    math_subscript(oMath, 'AUC', '0\u201312h')
    math_run(oMath, ' \u2264 ', italic=False)
    math_run(oMath, '60  mg\u00b7h/L', italic=False)


# ================================================================
# EQUATION BUILDERS - S8. PD MODEL
# ================================================================

def eq_impdh_inhibition(oMath):
    math_run(oMath, 'I')
    math_run(oMath, ' = ', italic=False)
    f, num, den = math_fraction(oMath)
    math_subscript(num, 'E', 'max')
    math_run(num, ' \u00b7 ')
    math_subsup(num, 'C', 'free', '\u03b3')
    math_subsup(den, 'IC', '50', '\u03b3', base_italic=False)
    math_run(den, ' + ')
    math_subsup(den, 'C', 'free', '\u03b3')

def eq_avg_impdh(oMath):
    math_subscript(oMath, 'I', 'avg')
    math_run(oMath, ' = ', italic=False)
    f, num, den = math_fraction(oMath)
    math_run(num, '1', italic=False)
    math_run(den, '\u03c4')
    body = math_nary(oMath, '\u222b', '0', '\u03c4')
    math_run(body, 'I')
    inner = math_delimiters(body)
    math_run(inner, 't')
    math_run(body, ' dt')

def eq_lymphocyte_ss(oMath):
    math_subscript(oMath, 'L', 'ss')
    math_run(oMath, ' = ', italic=False)
    math_subscript(oMath, 'L', '0')
    math_run(oMath, ' \u00b7 ', italic=False)
    f, num, den = math_fraction(oMath)
    math_subscript(num, 'k', 'prolif')
    math_run(num, ' \u00b7 ', italic=False)
    inner = math_delimiters(num)
    math_run(inner, '1 \u2212 ', italic=False)
    math_subscript(inner, 'I', 'avg')
    math_subscript(den, 'k', 'death')

def eq_rejection(oMath):
    math_subscript(oMath, 'P', 'rej')
    math_run(oMath, ' = ', italic=False)
    math_subscript(oMath, 'P', 'base')
    math_run(oMath, ' + ', italic=False)
    math_subscript(oMath, 'P', 'max')
    math_run(oMath, ' \u00b7 max', italic=False)
    inner = math_delimiters(oMath)
    math_subscript(inner, 'R', 'IMPDH')
    math_run(inner, ', ', italic=False)
    math_subscript(inner, 'R', 'AUC')

def eq_rejection_impdh(oMath):
    math_subscript(oMath, 'R', 'IMPDH')
    math_run(oMath, ' = ', italic=False)
    f, num, den = math_fraction(oMath)
    math_run(num, '1', italic=False)
    math_run(den, '1 + exp', italic=False)
    inner = math_delimiters(den)
    math_run(inner, '10 \u00b7 ', italic=False)
    inner2 = math_delimiters(inner)
    math_subscript(inner2, 'I', 'avg')
    math_run(inner2, ' \u2212 0.30', italic=False)

def eq_rejection_auc(oMath):
    math_subscript(oMath, 'R', 'AUC')
    math_run(oMath, ' = ', italic=False)
    f, num, den = math_fraction(oMath)
    math_run(num, '1', italic=False)
    math_run(den, '1 + exp', italic=False)
    inner = math_delimiters(den)
    math_run(inner, '0.15 \u00b7 ', italic=False)
    inner2 = math_delimiters(inner)
    math_run(inner2, 'AUC', italic=False)
    math_run(inner2, ' \u2212 30', italic=False)

def eq_gi_toxicity(oMath):
    math_subscript(oMath, 'P', 'GI')
    math_run(oMath, ' = ', italic=False)
    f, num, den = math_fraction(oMath)
    math_superscript(num, 'AUC', '\u03b3\u2081', base_italic=False)
    math_superscript(den, 'AUC', '\u03b3\u2081', base_italic=False)
    math_subscript(den, '', '50')
    math_run(den, ' + ')
    math_superscript(den, 'AUC', '\u03b3\u2081', base_italic=False)

def eq_leukopenia(oMath):
    math_subscript(oMath, 'P', 'leuko')
    math_run(oMath, ' = ', italic=False)
    f, num, den = math_fraction(oMath)
    math_subsup(num, 'I', 'avg', '\u03b3\u2082')
    math_subsup(den, 'IMPDH', '50', '\u03b3\u2082', base_italic=False)
    math_run(den, ' + ')
    math_subsup(den, 'I', 'avg', '\u03b3\u2082')

def eq_infection(oMath):
    math_subscript(oMath, 'P', 'inf')
    math_run(oMath, ' = ', italic=False)
    # Left brace with no right brace, wrapping a 3-row equation array (piecewise cases)
    inner = math_delimiters(oMath, open_char='{', close_char='')
    eqArr = math_sub(inner, 'eqArr')

    # Row 1: 0.05 if L_ss >= 1.0
    row1 = math_sub(eqArr, 'e')
    math_run(row1, '0.05', italic=False)
    math_run(row1, ',     if  ', italic=False)
    math_subscript(row1, 'L', 'ss')
    math_run(row1, ' \u2265 1.0', italic=False)

    # Row 2: 0.05 + 0.20 * (1.0 - L_ss) / (1.0 - L_thresh) if L_thresh <= L_ss < 1.0
    row2 = math_sub(eqArr, 'e')
    math_run(row2, '0.05 + 0.20 \u00b7 ', italic=False)
    f2, num2, den2 = math_fraction(row2)
    math_run(num2, '1.0 \u2212 ', italic=False)
    math_subscript(num2, 'L', 'ss')
    math_run(den2, '1.0 \u2212 ', italic=False)
    math_subscript(den2, 'L', 'thresh')
    math_run(row2, ',     if  ', italic=False)
    math_subscript(row2, 'L', 'thresh')
    math_run(row2, ' \u2264 ', italic=False)
    math_subscript(row2, 'L', 'ss')
    math_run(row2, ' < 1.0', italic=False)

    # Row 3: 0.25 + 0.35 * (L_thresh - L_ss) / L_thresh if L_ss < L_thresh
    row3 = math_sub(eqArr, 'e')
    math_run(row3, '0.25 + 0.35 \u00b7 ', italic=False)
    f3, num3, den3 = math_fraction(row3)
    math_subscript(num3, 'L', 'thresh')
    math_run(num3, ' \u2212 ', italic=False)
    math_subscript(num3, 'L', 'ss')
    math_subscript(den3, 'L', 'thresh')
    math_run(row3, ',     if  ', italic=False)
    math_subscript(row3, 'L', 'ss')
    math_run(row3, ' < ', italic=False)
    math_subscript(row3, 'L', 'thresh')

def eq_composite_adverse(oMath):
    math_subscript(oMath, 'P', 'adverse')
    math_run(oMath, ' = 1 \u2212 ', italic=False)
    inner = math_delimiters(oMath)
    math_run(inner, '1 \u2212 ', italic=False)
    math_subscript(inner, 'P', 'GI')
    math_run(oMath, ' \u00b7 ', italic=False)
    inner2 = math_delimiters(oMath)
    math_run(inner2, '1 \u2212 ', italic=False)
    math_subscript(inner2, 'P', 'leuko')
    math_run(oMath, ' \u00b7 ', italic=False)
    inner3 = math_delimiters(oMath)
    math_run(inner3, '1 \u2212 ', italic=False)
    math_subscript(inner3, 'P', 'inf')

def eq_therapeutic_index(oMath):
    math_run(oMath, 'TI', italic=False)
    math_run(oMath, ' = ', italic=False)
    math_subscript(oMath, 'I', 'avg')
    math_run(oMath, ' \u00d7 ', italic=False)
    inner = math_delimiters(oMath)
    math_run(inner, '1 \u2212 ', italic=False)
    math_subscript(inner, 'P', 'adverse')


# ================================================================
# EQUATION BUILDERS - S9. VALIDATION & STATISTICS
# ================================================================

def eq_gmfe(oMath):
    math_run(oMath, 'GMFE', italic=False)
    math_run(oMath, ' = exp', italic=False)
    inner = math_delimiters(oMath)
    math_run(inner, 'mean', italic=False)
    inner2 = math_delimiters(inner)
    inner3 = math_delimiters(inner2, '|', '|')
    math_run(inner3, 'log', italic=False)
    inner4 = math_delimiters(inner3)
    f, num, den = math_fraction(inner4)
    math_subscript(num, 'AUC', 'pred')
    math_subscript(den, 'AUC', 'obs')

def eq_fold_error(oMath):
    math_run(oMath, 'FE', italic=False)
    math_run(oMath, ' = ', italic=False)
    f, num, den = math_fraction(oMath)
    math_subscript(num, 'AUC', 'pred')
    math_subscript(den, 'AUC', 'obs')

def eq_cv(oMath):
    math_run(oMath, 'CV%', italic=False)
    math_run(oMath, ' = ', italic=False)
    f, num, den = math_fraction(oMath)
    math_run(num, 'SD')
    math_run(den, 'Mean')
    math_run(oMath, ' \u00d7 100', italic=False)


# ================================================================
# TABLE HELPERS
# ================================================================

def add_table_row(table, cells_data, bold=False, font_size=9):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        para = cell.paragraphs[0]
        run = para.add_run(str(text))
        run.font.size = Pt(font_size)
        run.font.name = 'Times New Roman'
        if bold:
            run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
    return row


SUPPLEMENTARY_FIGURES = [
    {
        "number": "S1",
        "filename": "FigureS1_Sensitivity.png",
        "title": "Factor decomposition sensitivity analysis",
        "legend": (
            "Figure S1. Factor decomposition sensitivity analysis. "
            "This figure decomposes the simulated exposure difference between the Western and Indian "
            "virtual transplant cohorts under the matched contemporary-practice scenario. Panel A "
            "shows the total MPA AUC_0-12h waterfall from the Western baseline to the Indian total, "
            "with the increment attributable to the combined body-weight, UGT-activity, and eGFR "
            "assumptions. Panel B shows the corresponding decomposition for free MPA AUC_0-12h. "
            "Panel C compares the body-weight distributions used to construct the Western and Indian "
            "virtual populations. Albumin and CNI co-therapy were deliberately matched between "
            "populations in this scenario, so they are not represented as between-population "
            "drivers in the waterfall."
        ),
    },
    {
        "number": "S2",
        "filename": "FigureS2_Profiles.png",
        "title": "Representative PK/PD profiles",
        "legend": (
            "Figure S2. Representative PK/PD profiles. "
            "Representative steady-state 12-hour dosing-interval simulations are shown for a Western "
            "reference patient receiving 1000 mg MMF BID, an Indian reference patient receiving "
            "1000 mg MMF BID, the same Indian reference patient receiving 750 mg BID, and the same "
            "Indian reference patient receiving the weight-based 12 mg/kg BID strategy rounded to an "
            "available tablet strength. Panel A shows total MPA concentration-time profiles, Panel B "
            "shows free MPA concentration-time profiles, and Panel C shows the corresponding IMPDH-II "
            "inhibition profiles. Panel D summarizes model-derived clinical outcome metrics, including "
            "rejection probability, GI toxicity probability, and the therapeutic index, illustrating "
            "how weight-based dose reduction lowers toxicity while preserving pharmacodynamic activity."
        ),
    },
    {
        "number": "S3",
        "filename": "FigureS3_PD_Heatmaps.png",
        "title": "PD outcome heatmaps by weight and AUC",
        "legend": (
            "Figure S3. PD outcome heatmaps by body weight and total MPA exposure. "
            "Each point represents a virtual patient plotted by body weight and total MPA AUC_0-12h, "
            "with color indicating the predicted pharmacodynamic outcome. The three panels show GI "
            "toxicity probability, leukopenia probability, and therapeutic index, respectively. "
            "Horizontal dashed lines mark the conventional therapeutic AUC range boundaries of "
            "30 and 60 mg.h/L. The plots show that low body weight and higher total exposure shift "
            "patients toward higher toxicity probabilities, whereas the therapeutic-index panel "
            "summarizes the balance between immunosuppressive efficacy and adverse-event risk."
        ),
    },
    {
        "number": "S4",
        "filename": "FigureS4_Variability_Decomposition.png",
        "title": "Variability decomposition analysis",
        "legend": (
            "Figure S4. Variability decomposition analysis. "
            "This figure quantifies sources of interindividual variability in simulated total MPA "
            "AUC_0-12h for Western and Indian virtual populations. Panel A shows univariate Pearson "
            "correlations between AUC and model covariates. Panel B shows partial correlations after "
            "adjustment for the other covariates. Panel C ranks Indian-population variance attribution "
            "by the reduction in AUC CV% obtained when each covariate is fixed at its population mean. "
            "Panel D compares the same variance-attribution metric between Western and Indian "
            "populations. Panel E overlays simulated weight-AUC relationships with the theoretical "
            "AUC proportional to BW^-0.75 relationship, and Panel F shows the derivative of that "
            "relationship, demonstrating the steeper exposure sensitivity at the Indian mean body "
            "weight compared with the Western mean body weight."
        ),
    },
    {
        "number": "S5",
        "filename": "FigureS5_ParameterSensitivity.png",
        "title": "Parameter uncertainty sensitivity analysis",
        "legend": (
            "Figure S5. Parameter uncertainty sensitivity analysis. "
            "One-at-a-time local sensitivity analysis of total MPA AUC_0-12h was performed at the "
            "Indian mean-patient baseline receiving 1 g MMF BID. Six structural PBPK parameters were "
            "perturbed by +/-20% from their nominal values: reference free fraction at albumin 4.0 g/dL, "
            "UGT1A9 intrinsic clearance, UGT2B7 intrinsic clearance, central volume, absorption rate, "
            "and biliary efflux rate. Bars show the percentage change in total AUC relative to the "
            "baseline simulation, with parameters ordered by effect size. The analysis identifies "
            "the reference free fraction and UGT1A9 intrinsic clearance as the dominant local "
            "uncertainty drivers while showing that single-parameter perturbations remain smaller "
            "than the population-level Indian-vs-Western exposure difference."
        ),
    },
]


def add_supplementary_figures(doc):
    """Append supplementary figures in fixed numeric order with detailed legends."""
    doc.add_page_break()
    doc.add_heading('Supplementary Figures', level=1)

    for idx, fig in enumerate(SUPPLEMENTARY_FIGURES):
        if idx > 0:
            doc.add_page_break()

        doc.add_heading(f"Figure {fig['number']}. {fig['title']}", level=2)
        fig_path = os.path.join(FIG_DIR, fig["filename"])
        if os.path.exists(fig_path):
            picture = doc.add_picture(fig_path, width=Inches(6.5))
            picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            missing = doc.add_paragraph()
            run = missing.add_run(f"[Missing figure file: {fig['filename']}]")
            run.bold = True
            run.font.color.rgb = RGBColor(192, 0, 0)

        caption = doc.add_paragraph()
        caption.paragraph_format.space_before = Pt(6)
        caption.paragraph_format.space_after = Pt(12)
        run = caption.add_run(fig["legend"])
        run.font.size = Pt(10)
        run.italic = True
        run.font.name = 'Times New Roman'


# ================================================================
# DOCUMENT BUILDER
# ================================================================

def create_supplement():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)

    # ================================================================
    # TITLE PAGE
    # ================================================================
    for _ in range(3):
        doc.add_paragraph('')

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Supplementary Material: Mathematical Framework')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'

    doc.add_paragraph('')
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        'Complete Mathematical Specification of the Semi-Mechanistic PBPK/PD Model\n'
        'for Mycophenolic Acid in Indian and Western Transplant Populations'
    )
    run.italic = True
    run.font.size = Pt(12)

    doc.add_paragraph('')
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        'This supplement provides the complete mathematical specification for all equations '
        'referenced in the main manuscript. Equations are numbered S1\u2013S33 with cross-references '
        'to the corresponding manuscript equations where applicable.'
    )
    run.font.size = Pt(10)

    doc.add_page_break()

    # ================================================================
    # TABLE OF CONTENTS
    # ================================================================
    toc = doc.add_heading('Contents', level=1)
    toc.runs[0].font.size = Pt(14)

    toc_items = [
        'S1.  Hepatic Clearance Model (Equations S1\u2013S4)',
        'S2.  Protein Binding and Free Fraction (Equations S5\u2013S7)',
        'S3.  Allometric and Physiological Scaling (Equations S8\u2013S14)',
        'S4.  Ordinary Differential Equation System (Equations S15a\u2013S15e)',
        'S5.  ODE Rate Expressions (Equations S16a\u2013S16e)',
        'S6.  Drug\u2013Drug Interactions (Equations S17a\u2013S17b)',
        'S7.  Dose Conversion and Exposure Metrics (Equations S18\u2013S21)',
        'S8.  Pharmacodynamic Model (Equations S22\u2013S30)',
        'S9.  Validation and Statistical Metrics (Equations S31\u2013S33)',
        'S10. Parameter Summary Tables',
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ================================================================
    # S1. HEPATIC CLEARANCE MODEL
    # ================================================================
    doc.add_heading('S1. Hepatic Clearance Model', level=1)

    doc.add_paragraph(
        'MPA hepatic metabolism is represented as a UGT-mediated process dominated by UGT1A9, '
        'with a UGT2B7 contribution to acyl-glucuronide formation. In the model, total intrinsic '
        'clearance is partitioned as 75% UGT1A9 and 25% UGT2B7 for interpretability and sensitivity '
        'analysis. Hepatic clearance is modeled using the well-stirred model, which assumes rapid '
        'equilibration between blood and hepatocyte compartments.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Equation S1 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('(Manuscript Eq. 1) \u2014 Well-stirred hepatic clearance:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_well_stirred, number='S1')

    doc.add_paragraph(
        'where CL\u2095 is hepatic clearance (L/h), Q\u2095 is hepatic blood flow (L/h), '
        'f\u1d64 is the unbound fraction of MPA in plasma, and CL\u1d62\u2099\u209c is the total '
        'intrinsic clearance (L/h). The well-stirred model is the standard approach for drugs '
        'with flow-independent clearance (low extraction ratio).'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S2 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('(Manuscript Eq. 2) \u2014 Hepatic extraction ratio:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_extraction_ratio, number='S2')

    doc.add_paragraph(
        'For MPA, E \u2248 0.13 (low extraction ratio drug). This has important pharmacokinetic '
        'consequences: changes in protein binding directly affect total drug clearance, but '
        'not unbound drug clearance. Numerically: f\u1d64 \u00d7 CL\u1d62\u2099\u209c = '
        '0.03 \u00d7 460 = 13.8 L/h, while Q\u2095 \u2248 90 L/h, yielding '
        'CL\u2095 = 90 \u00d7 13.8 / (90 + 13.8) = 12.0 L/h.'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S3 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('(Manuscript Eq. 3) \u2014 Low extraction ratio approximation:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_low_e_approx, number='S3')

    doc.add_paragraph(
        'This approximation holds when f\u1d64 \u00b7 CL\u1d62\u2099\u209c \u226a Q\u2095. '
        'It explains the "protein binding paradox": lower albumin increases f\u1d64, which '
        'increases CL\u2095 proportionally. The net effect on total AUC is attenuated because '
        'both clearance and free fraction change in the same direction. However, free MPA '
        'concentrations (and thus pharmacodynamic effect) are genuinely increased.'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S4 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 Total intrinsic clearance as the sum of UGT pathway contributions:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_total_clint, number='S4')

    doc.add_paragraph(
        'Reference values: CL\u1d62\u2099\u209c,UGT1A9 = 345 L/h, CL\u1d62\u2099\u209c,UGT2B7 = 115 L/h '
        '(total = 460 L/h). The fraction metabolized by each pathway '
        '(f\u2098,UGT1A9 = 0.75, f\u2098,UGT2B7 = 0.25) is the ratio of each contribution to the '
        'total and is reported for interpretability only; it is not applied as a coefficient in S4. '
        'Individual-specific UGT activity multipliers, sampled from log-normal distributions '
        '(Western UGT1A9 1.00 \u00b1 0.15, UGT2B7 1.00 \u00b1 0.15; Indian UGT1A9 0.95 \u00b1 0.18, '
        'UGT2B7 0.92 \u00b1 0.20), are applied separately during patient-level scaling along with '
        'liver weight (Equation S10).'
    )

    # ================================================================
    # S2. PROTEIN BINDING
    # ================================================================
    doc.add_page_break()
    doc.add_heading('S2. Protein Binding and Free Fraction', level=1)

    doc.add_paragraph(
        'MPA is approximately 97\u201398% bound to serum albumin. The unbound fraction is modulated by '
        'two physiological factors: (1) albumin concentration and (2) competitive displacement by the '
        'glucuronide metabolite MPAG. The free fraction is dynamically recalculated at each ODE time step.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Equation S5 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('(Manuscript Eq. 4) \u2014 Albumin-dependent baseline free fraction:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_fu_albumin, number='S5')

    doc.add_paragraph(
        'where f\u1d64,ref = 0.03 at Alb\u1d63\u1d49\u1da0 = 4.0 g/dL. The exponent 1.2 '
        'reflects the nonlinear albumin\u2013binding relationship for highly bound acidic drugs '
        'used in this model. In this study, albumin was matched at 4.0 \u00b1 0.4 g/dL '
        'for both populations as a model-design assumption, removing it as a confounding variable.'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S6 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 MPAG displacement factor (E\u2098\u2090\u2093-type saturable kinetics):')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_mpag_displacement, number='S6')

    doc.add_paragraph(
        'MPAG competitively displaces MPA from albumin binding sites (de Winter et al. 2009). '
        'The maximum displacement is 50% increase in f\u1d64 at saturating MPAG concentrations, '
        'with half-maximal effect at C\u2098\u209a\u2090\u1d4d = 30 mg/L. This effect is clinically '
        'significant during cyclosporine co-therapy, where OATP1B inhibition raises systemic MPAG.'
    ).add_run(f' {cite(7)}')

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S7 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 Combined free fraction with physiological bounds:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_fu_combined, number='S7')

    doc.add_paragraph(
        'The combined free fraction is bounded between 0.005 (minimum physiologically feasible) '
        'and 0.30 (maximum displacement scenario) to prevent numerical instabilities in the ODE solver.'
    )

    # ================================================================
    # S3. ALLOMETRIC SCALING
    # ================================================================
    doc.add_page_break()
    doc.add_heading('S3. Allometric and Physiological Scaling', level=1)

    doc.add_paragraph(
        'Patient-specific PK parameters are derived from reference values (70 kg Western adult) '
        'using standard allometric scaling principles. These relationships are central to the '
        'body weight\u2013exposure mechanism that drives the population differences in this study.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Equation S8 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('(Manuscript Eq. 5a) \u2014 Volume scaling (exponent 1.0):')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_allometric_volume, number='S8')

    doc.add_paragraph(
        'Distribution volumes (V\u1d9c, V\u209a, V\u2098\u209a\u2090\u1d4d) scale linearly with '
        'body weight, reflecting proportional increases in tissue mass and body water. '
        'Reference values at 70 kg: V\u1d9c = 50 L, V\u209a = 150 L, V\u2098\u209a\u2090\u1d4d = 15 L.'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S9 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('(Manuscript Eq. 5b) \u2014 Systemic clearance scaling (exponent 0.75):')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_allometric_clearance, number='S9')

    doc.add_paragraph(
        'Intercompartmental clearance (Q) and the reference hepatic blood flow (Q\u2095, see S11) scale '
        'with the standard metabolic allometric exponent of 0.75 (West et al. 1997). Reference: '
        'Q = 30 L/h at 70 kg. Intrinsic clearance (CL\u1d62\u2099\u209c) is not scaled directly by '
        '(BW/70)\u2070\u00b7\u2077\u2075; it is scaled by relative liver weight (Equation S10), so that '
        'CL\u1d62\u2099\u209c,\u1d62 = CL\u1d62\u2099\u209c,\u1d63\u2091\u0192 \u00b7 (LW\u1d62 / LW\u1d63\u2091\u0192) \u00b7 A_UGT, '
        'i.e. CL\u1d62\u2099\u209c \u221d BW\u2070\u00b7\u2078\u2076 before pharmacogenomic activity multipliers A_UGT are applied.'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S10 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 Liver weight allometric scaling:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_liver_weight, number='S10')

    doc.add_paragraph(
        'Liver weight scales with exponent 0.86 (Johnson et al. 2005), intermediate between '
        'the volume (1.0) and clearance (0.75) exponents. Intrinsic clearance is proportional '
        'to liver weight, reflecting total hepatic UGT enzyme content. For a 58 kg patient: '
        'LW = 1.5 \u00d7 (58/70)\u2070\u22c8\u2078\u2076 = 1.28 kg (vs 1.50 kg at 70 kg).'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S11 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 Hepatic blood flow scaling:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_hepatic_blood_flow, number='S11')

    doc.add_paragraph(
        'Hepatic blood flow scales with the 0.75 exponent. The reference value of '
        '90 L/h corresponds to approximately 25% of cardiac output for a 70 kg adult.'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S12 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 Du Bois body surface area:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_bsa, number='S12')

    doc.add_paragraph(
        'where H is height (cm) and BW is body weight (kg). BSA is used for anthropometric '
        'characterization but does not directly enter PK calculations in this model.'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S13 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 GFR-proportional renal clearance of MPAG:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_renal_clearance, number='S13')

    doc.add_paragraph(
        'MPAG is eliminated renally with clearance proportional to GFR. '
        'Reference: CL\u1d63,ref = 7.5 L/h at GFR\u1d63\u1d49\u1da0 = 60 mL/min. Reduced GFR in '
        'transplant patients leads to MPAG accumulation, which can increase MPA f\u1d64 via '
        'the displacement mechanism (Equation S6).'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S14 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 Gut wall bioavailability (UGT1A8/UGT1A9-dependent):')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_gut_wall, number='S14')

    doc.add_paragraph(
        'Presystemic gut wall metabolism is mediated by UGT1A8 (closely related to and co-regulated '
        'with UGT1A9). Higher UGT1A9 activity reduces oral bioavailability. F\u1d4d\u2081,ref = 0.90. '
        'The max() function prevents division by very small activity values, with a floor of 0.5.'
    )

    # ================================================================
    # S4. ODE SYSTEM
    # ================================================================
    doc.add_page_break()
    doc.add_heading('S4. Ordinary Differential Equation System', level=1)

    doc.add_paragraph(
        'The PBPK model is defined by a system of six coupled ordinary differential equations '
        'representing the mass balance in each physiological compartment.'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('State variables:')
    run.bold = True
    run.font.size = Pt(10)
    doc.add_paragraph(
        'y[0] = A\u2091\u1d64\u209c : Amount in gut absorption depot (mg)\n'
        'y[1] = C\u2098\u209a\u2090 : MPA concentration in central compartment (mg/L)\n'
        'y[2] = A\u209a : Amount in peripheral compartment (mg)\n'
        'y[3] = C\u2098\u209a\u2090\u1d4d : MPAG plasma concentration (mg/L)\n'
        'y[4] = A\u2091\u1d62\u2097\u1d49 : MPAG amount in bile/gut pool for EHC (mg)\n'
        'y[5] = A\u2090\u2091\u209b : Cumulative amount absorbed (mg, for tracking)'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S15a ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 Gut absorption depot:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_ode_gut, number='S15a')

    doc.add_paragraph(
        'First-order absorption with rate constant k\u2090 = 4.0 /h.'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S15b ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('(Manuscript Eq. 6) \u2014 Central compartment MPA concentration:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_ode_central, number='S15b')

    doc.add_paragraph(
        'This is the core mass balance equation. The rate terms (R) are defined in Section S5.'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S15c ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 Peripheral compartment:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_ode_periph, number='S15c')

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S15d ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 MPAG plasma concentration:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_ode_mpag, number='S15d')

    doc.add_paragraph(
        'MPAG is formed from hepatic MPA metabolism, eliminated renally, and excreted into bile '
        'via OATP1B-dependent hepatic uptake followed by ABCC2-mediated canalicular excretion.'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S15e ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 Biliary MPAG pool (enterohepatic recirculation):')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_ode_bile, number='S15e')

    doc.add_paragraph(
        'The system is initialized with the MPA-equivalent dose in the gut compartment and all other '
        'states at zero (or carried forward from the previous dosing interval for steady-state simulations). '
        'The adaptive-step LSODA integrator (rtol = 10\u207b\u2078, atol = 10\u207b\u00b9\u2070) ensures '
        'numerical stability. Steady state is reached after 14 BID doses (7 days).'
    )

    # ================================================================
    # S5. RATE EXPRESSIONS
    # ================================================================
    doc.add_page_break()
    doc.add_heading('S5. ODE Rate Expressions', level=1)

    doc.add_paragraph(
        'The individual rate terms appearing in Equations S15a\u2013S15e are defined below.'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S16a ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 Absorption rate:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_rate_absorption, number='S16a')

    doc.add_paragraph(
        'where k\u2090 = 4.0 /h, F\u2092\u1d63\u2090\u2097 = 0.81, and F\u1d4d\u2081 is '
        'the gut wall bioavailability (Equation S14). A lag time of t\u2097\u2090\u1d4d = 0.25 h '
        'accounts for gastric emptying.'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S16b ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 Intercompartmental distribution rates:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_rate_distribution, number='S16b')

    doc.add_paragraph(
        'Q = 30 L/h at 70 kg reference, scaled allometrically (Equation S9).'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S16c ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 Hepatic metabolism rate:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_rate_metabolism, number='S16c')

    doc.add_paragraph(
        'CL\u2095 is computed dynamically at each time step via Equation S1 using the current '
        'f\u1d64, which depends on the instantaneous MPAG concentration (Equations S5\u2013S7).'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S16d ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 Enterohepatic recirculation return:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_rate_ehc, number='S16d')

    doc.add_paragraph(
        'k\u2091\u1d64\u209c = 0.15 /h (gut bacterial \u03b2-glucuronidase deconjugation rate), '
        'f\u1d31\u1d34\u1d30 = 0.40. MPAG in bile is deconjugated to regenerate MPA, producing '
        'the characteristic secondary plasma peak at 6\u20128 hours post-dose.'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S16e ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 OATP1B-dependent biliary excretion rate:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_rate_biliary, number='S16e')

    doc.add_paragraph(
        'k\u2091\u1d62\u2097\u1d49 = 0.8 /h (ABCC2/MRP2-mediated canalicular efflux rate). '
        'Biliary excretion requires prior hepatic uptake via OATP1B1/1B3 transporters, modulated '
        'by f\u2092\u2090\u209c\u209a\u2081\u2091 (= 1.0 with tacrolimus, reduced with cyclosporine).'
    )

    # ================================================================
    # S6. DRUG INTERACTIONS
    # ================================================================
    doc.add_page_break()
    doc.add_heading('S6. Drug\u2013Drug Interactions: Cyclosporine\u2013MPA', level=1)

    doc.add_paragraph(
        'Cyclosporine (CsA) was represented through two complementary transporter-mediated scenario '
        'mechanisms. These interactions are clinically significant but were minimized in this study '
        'by matching tacrolimus co-therapy at ~93% in both populations (approximately 7% on cyclosporine).'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S17a ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 CsA inhibition of ABCC2-mediated biliary excretion:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_csa_ehc, number='S17a')

    doc.add_paragraph(
        'I\u2090\u2091\u1d9c\u1d9c\u2082 = 0.40 (40% reduction in EHC). CsA inhibits ABCC2/MRP2 '
        'on the canalicular membrane, reducing biliary MPAG excretion and consequently the fraction '
        'available for enterohepatic recirculation. The 40% value is a model scenario coefficient, '
        'not a directly estimated population effect.'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S17b ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 CsA inhibition of OATP1B1/1B3 hepatic uptake:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_csa_oatp, number='S17b')

    doc.add_paragraph(
        'I\u2092\u2090\u209c\u209a = 0.60 (60% reduction in hepatic uptake). CsA inhibits OATP1B1 '
        '(K\u1d62 \u2248 0.2\u20132 \u00b5M), reducing hepatic MPAG uptake and raising systemic MPAG. '
        'Elevated MPAG secondarily increases MPA f\u1d64 via the displacement mechanism (Equation S6). '
        'The 60% value is a model scenario coefficient used to represent reduced MPAG hepatic uptake, '
        'not a directly estimated population effect.'
    )

    # ================================================================
    # S7. DOSE & EXPOSURE
    # ================================================================
    doc.add_page_break()
    doc.add_heading('S7. Dose Conversion and Exposure Metrics', level=1)

    p = doc.add_paragraph()
    run = p.add_run('Equation S18 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 MPA-equivalent dose from MMF prodrug:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_mpa_dose, number='S18')

    doc.add_paragraph(
        'MW\u2098\u209a\u2090 = 320.3 g/mol, MW\u2098\u2098\u1da0 = 433.5 g/mol, '
        'F\u1d9c\u2092\u2099\u1d65 = 1.0 (complete presystemic ester hydrolysis). '
        'Thus 1000 mg MMF yields 738.7 mg MPA equivalent entering the absorption depot.'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S19 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 Area under the concentration\u2013time curve:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_auc_trapezoidal, number='S19')

    doc.add_paragraph(
        'AUC\u2080\u208b\u03c4 is computed numerically using the trapezoidal rule over the '
        'steady-state dosing interval (\u03c4 = 12 hours for BID dosing). The time resolution '
        'is 0.05 hours (3 minutes), ensuring accurate numerical integration.'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S20 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 Free (unbound) MPA AUC:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_auc_free, number='S20')

    doc.add_paragraph(
        'Free MPA AUC is the pharmacologically relevant exposure metric, representing the '
        'concentration available to interact with the IMPDH-II target. Because f\u1d64 varies '
        'dynamically with MPAG concentration across the dosing interval, '
        'AUC\u1da0\u1d63\u1d49\u1d49 \u2260 f\u0305\u1d64 \u00d7 AUC\u209c\u2092\u209c\u2090\u2097.'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S21 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 Therapeutic AUC window:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_auc_target, number='S21')

    doc.add_paragraph(
        'The consensus therapeutic window for total MPA AUC\u2080\u208b\u2081\u2082\u2095 in renal '
        'transplant recipients is 30\u201360 mg\u00b7h/L, consistent with MPA TDM literature and '
        'Indian renal-transplant exposure studies. '
        'Below 30 mg\u00b7h/L, rejection risk rises; above 60 mg\u00b7h/L, toxicity risk increases.'
    ).add_run(f' {cite(1, 3, 5, 6)}')

    # ================================================================
    # S8. PD MODEL
    # ================================================================
    doc.add_page_break()
    doc.add_heading('S8. Pharmacodynamic Model', level=1)

    doc.add_paragraph(
        'The PD model links free MPA concentrations to IMPDH-II inhibition, downstream lymphocyte '
        'dynamics, and clinical outcome probabilities. All PD parameters are identical between '
        'Western and Indian populations \u2014 all PD differences arise exclusively from upstream PK.'
    )

    doc.add_heading('S8.1 IMPDH-II Inhibition', level=2)

    p = doc.add_paragraph()
    run = p.add_run('Equation S22 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('(Manuscript Eq. 7) \u2014 Sigmoidal E\u2098\u2090\u2093 IMPDH-II inhibition:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_impdh_inhibition, number='S22')

    doc.add_paragraph(
        'E\u2098\u2090\u2093 = 1.0 (complete inhibition possible), IC\u2085\u2080 = 0.15 mg/L '
        '(free MPA), \u03b3 = 1.5 (Hill coefficient, mild cooperativity). These are phenomenological '
        'calibration parameters informed by published IMPDH activity data, '
        'targeting ~50% IMPDH suppression at the midpoint of the therapeutic free MPA range.'
    ).add_run(f' {cite(15, 16)}')

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S23 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 Time-averaged IMPDH inhibition:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_avg_impdh, number='S23')

    doc.add_paragraph(
        'I\u2090\u1d65\u1d4d integrates the fluctuating IMPDH response over the 12-hour '
        'dosing interval (\u03c4), providing a single PD summary metric computed by trapezoidal rule.'
    )

    doc.add_heading('S8.2 Lymphocyte Dynamics', level=2)

    p = doc.add_paragraph()
    run = p.add_run('Equation S24 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 Steady-state lymphocyte count:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_lymphocyte_ss, number='S24')

    doc.add_paragraph(
        'L\u2080 = 2.0 \u00d7 10\u2079/L (baseline lymphocyte count), '
        'k\u209a\u1d63\u2092\u2097\u1d62\u1da0 = k\u1d48\u1d49\u2090\u209c\u2095 = 0.05 /h '
        '(balanced at baseline so L\u209b\u209b = L\u2080 when I = 0). IMPDH inhibition reduces '
        'the effective proliferation rate. L\u209b\u209b is bounded between 0.2 and 3.0 \u00d7 10\u2079/L.'
    )

    doc.add_heading('S8.3 Clinical Outcome Probabilities', level=2)

    p = doc.add_paragraph()
    run = p.add_run('Equation S25 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 Rejection probability:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_rejection, number='S25')

    doc.add_paragraph(
        'P\u2091\u2090\u209b\u1d49 = 0.10, P\u2098\u2090\u2093 = 0.30. The composite risk is '
        'driven by the maximum of two logistic components:'
    )

    p = doc.add_paragraph()
    run = p.add_run('Equation S25a ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 IMPDH-based rejection risk:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_rejection_impdh, number='S25a')

    doc.add_paragraph(
        'Steep logistic (slope = 10) centered at I\u2090\u1d65\u1d4d = 0.30. Below this threshold, '
        'rejection risk rises sharply.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Equation S25b ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 AUC-based rejection risk:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_rejection_auc, number='S25b')

    doc.add_paragraph(
        'Gentler logistic (slope = 0.15) centered at AUC = 30 mg\u00b7h/L.'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S26 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 GI toxicity (E\u2098\u2090\u2093 model):')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_gi_toxicity, number='S26')

    doc.add_paragraph(
        'AUC\u2085\u2080 = 70 mg\u00b7h/L, \u03b3\u2081 = 3.0 (steep dose\u2013response). '
        'GI adverse events (diarrhea, nausea) are the most common MPA toxicity.'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S27 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 Leukopenia probability:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_leukopenia, number='S27')

    doc.add_paragraph(
        'IMPDH\u2085\u2080 = 0.65, \u03b3\u2082 = 4.0. Leukopenia reflects excessive suppression '
        'of lymphocyte and granulocyte proliferation.'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S28 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 Infection risk (piecewise lymphocyte-dependent):')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_infection, number='S28')

    doc.add_paragraph(
        'Infection risk is a piecewise function of L\u209b\u209b: 5% baseline when '
        'L\u209b\u209b \u2265 1.0 \u00d7 10\u2079/L; linearly increasing to 25% at the '
        'lymphopenia threshold (0.5 \u00d7 10\u2079/L); steeply rising below '
        '(up to 60% at the minimum viable count of 0.2 \u00d7 10\u2079/L).'
    )

    doc.add_heading('S8.4 Composite Metrics', level=2)

    p = doc.add_paragraph()
    run = p.add_run('Equation S29 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 Composite adverse event probability (independence assumption):')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_composite_adverse, number='S29')

    doc.add_paragraph(
        'Individual adverse event probabilities are combined assuming statistical independence. '
        'This provides a conservative (lower bound) estimate of composite toxicity risk.'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S30 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('(Manuscript Eq. 8) \u2014 Therapeutic index:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_therapeutic_index, number='S30')

    doc.add_paragraph(
        'TI captures the efficacy\u2013safety balance: I\u2090\u1d65\u1d4d reflects immunosuppressive '
        'efficacy (higher = better protection against rejection) and (1 \u2212 P\u2090\u1d48\u1d65\u1d49\u1d63\u209b\u1d49) '
        'reflects safety (higher = fewer adverse events). TI ranges from 0 to ~1. A value \u2265 0.20 '
        'is considered a favorable balance. In this study: Western 1g BID = 0.215, '
        'Indian 1g BID = 0.156, Indian 12 mg/kg = 0.219.'
    )

    # ================================================================
    # S9. VALIDATION & STATISTICS
    # ================================================================
    doc.add_page_break()
    doc.add_heading('S9. Validation and Statistical Metrics', level=1)

    p = doc.add_paragraph()
    run = p.add_run('Equation S31 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('(Manuscript Eq. 10) \u2014 Geometric Mean Fold Error:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_gmfe, number='S31')

    doc.add_paragraph(
        'GMFE is the standard PBPK validation metric (Abduljalil et al. 2014). The absolute value '
        'of the log-ratio ensures symmetric penalization of over- and under-predictions. '
        'Regulatory acceptance: GMFE < 2.0. Our model: GMFE = 1.23 across 5 published studies.'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S32 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('\u2014 Individual fold error:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_fold_error, number='S32')

    doc.add_paragraph(
        'FE = 1.0 is ideal. Stringent criterion: 0.80 \u2264 FE \u2264 1.25 (bioequivalence-like). '
        'Acceptable: 0.50 \u2264 FE \u2264 2.00 (within 2-fold). In this study, 3/5 studies met the '
        'stringent criterion and 5/5 met the acceptable criterion.'
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Equation S33 ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run('(Manuscript Eq. 9) \u2014 Coefficient of variation:')
    run.font.size = Pt(10)
    create_display_equation(doc, eq_cv, number='S33')

    doc.add_paragraph(
        'CV% quantifies interindividual variability relative to the mean. Indian patients '
        'showed higher AUC variability (CV% 36.2% vs 30.3%) attributable to wider UGT enzyme '
        'activity distributions and the nonlinear weight\u2013exposure relationship.'
    )

    # ================================================================
    # S10. PARAMETER SUMMARY TABLES
    # ================================================================
    doc.add_page_break()
    doc.add_heading('S10. Parameter Summary Tables', level=1)

    doc.add_paragraph(
        'Tables S-A through S-C provide complete parameter listings for the PBPK/PD model.'
    )

    # Table S-A: PK Parameters
    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Table S-A. Pharmacokinetic Model Parameters (Reference: 70 kg Western Adult)')
    run.bold = True
    run.font.size = Pt(10)

    pk_table = doc.add_table(rows=1, cols=4)
    pk_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    pk_table.style = 'Light Grid Accent 1'
    for i, text in enumerate(['Parameter', 'Symbol', 'Value', 'Source']):
        cell = pk_table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    pk_params = [
        ('Absorption rate constant', 'k\u2090', '4.0 /h', cite(1)),
        ('Absorption lag time', 't\u2097\u2090\u1d4d', '0.25 h', cite(1)),
        ('Oral bioavailability', 'F\u2092\u1d63\u2090\u2097', '0.81', 'Model calibration'),
        ('Gut wall bioavailability', 'F\u1d4d\u2081', '0.90', 'Model calibration'),
        ('Central volume', 'V\u1d9c', '50 L', cite(7)),
        ('Peripheral volume', 'V\u209a', '150 L', cite(7)),
        ('Intercompartmental CL', 'Q', '30 L/h', cite(7)),
        ('UGT1A9 intrinsic CL', 'CL\u1d62\u2099\u209c,1A9', '345 L/h', 'Model calibration'),
        ('UGT2B7 intrinsic CL', 'CL\u1d62\u2099\u209c,2B7', '115 L/h', 'Model calibration'),
        ('UGT1A9 fraction', 'f\u2098,1A9', '0.75', f'Model split based on {cite(12)}'),
        ('UGT2B7 fraction', 'f\u2098,2B7', '0.25', f'Model split based on {cite(12)}'),
        ('Biliary excretion rate', 'k\u2091\u1d62\u2097\u1d49', '0.8 /h', cite(8)),
        ('EHC release rate', 'k\u2091\u1d64\u209c', '0.15 /h', cite(8)),
        ('EHC fraction', 'f\u1d31\u1d34\u1d30', '0.40', cite(8)),
        ('MPAG central volume', 'V\u2098\u209a\u2090\u1d4d', '15 L', cite(7)),
        ('MPAG renal clearance', 'CL\u1d63,MPAG', '7.5 L/h', 'Model calibration'),
        ('Reference free fraction', 'f\u1d64,ref', '0.03', 'Model calibration'),
        ('Reference albumin', 'Alb\u1d63\u1d49\u1da0', '4.0 g/dL', 'Matched scenario assumption'),
        ('Albumin\u2013binding exponent', 'n', '1.2', cite(7)),
        ('MPAG displacement E\u2098\u2090\u2093', '\u2014', '0.50', cite(7)),
        ('MPAG displacement EC\u2085\u2080', '\u2014', '30 mg/L', cite(7)),
        ('CsA ABCC2 inhibition', 'I\u2090\u2091\u1d9c\u1d9c\u2082', '0.40', 'Scenario coefficient'),
        ('CsA OATP1B inhibition', 'I\u2092\u2090\u209c\u209a', '0.60', 'Scenario coefficient'),
        ('MMF molecular weight', 'MW\u2098\u2098\u1da0', '433.5 g/mol', 'Standard'),
        ('MPA molecular weight', 'MW\u2098\u209a\u2090', '320.3 g/mol', 'Standard'),
        ('MPAG molecular weight', 'MW\u2098\u209a\u2090\u1d4d', '496.5 g/mol', 'Standard'),
    ]
    for row_data in pk_params:
        add_table_row(pk_table, row_data, font_size=8)

    # Table S-B: PD Parameters
    doc.add_paragraph('')
    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Table S-B. Pharmacodynamic Model Parameters (Identical for Both Populations)')
    run.bold = True
    run.font.size = Pt(10)

    pd_table = doc.add_table(rows=1, cols=4)
    pd_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    pd_table.style = 'Light Grid Accent 1'
    for i, text in enumerate(['Parameter', 'Symbol', 'Value', 'Source']):
        cell = pd_table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    pd_params = [
        ('IMPDH E\u2098\u2090\u2093', 'E\u2098\u2090\u2093', '1.0', 'Theoretical maximum'),
        ('IMPDH IC\u2085\u2080 (free MPA)', 'IC\u2085\u2080', '0.15 mg/L', 'Phenomenological calibration'),
        ('Hill coefficient', '\u03b3', '1.5', 'Phenomenological calibration'),
        ('Baseline lymphocyte count', 'L\u2080', '2.0 \u00d7 10\u2079/L', 'Standard'),
        ('Proliferation rate', 'k\u209a\u1d63\u2092\u2097\u1d62\u1da0', '0.05 /h', 'Model calibration'),
        ('Death rate', 'k\u1d48\u1d49\u2090\u209c\u2095', '0.05 /h', 'Balanced at steady state'),
        ('GI toxicity AUC\u2085\u2080', 'AUC\u2085\u2080,GI', '70 mg\u00b7h/L', 'Phenomenological calibration'),
        ('GI toxicity Hill coefficient', '\u03b3\u2081', '3.0', 'Model calibration'),
        ('Leukopenia IMPDH\u2085\u2080', 'IMPDH\u2085\u2080', '0.65', 'Phenomenological calibration'),
        ('Leukopenia Hill coefficient', '\u03b3\u2082', '4.0', 'Model calibration'),
        ('Infection lymphocyte threshold', 'L\u209c\u2095\u1d63\u1d49\u209b\u2095', '0.5 \u00d7 10\u2079/L', 'Standard'),
        ('Baseline infection risk', '\u2014', '0.05', 'Clinical estimate'),
        ('Maximum infection risk', '\u2014', '0.60', 'Clinical estimate'),
        ('Baseline rejection probability', 'P\u2091\u2090\u209b\u1d49', '0.10', 'Clinical estimate'),
        ('Max additional rejection risk', 'P\u2098\u2090\u2093', '0.30', 'Phenomenological calibration'),
        ('Rejection IMPDH threshold', '\u2014', '0.30', 'Phenomenological calibration'),
        ('Rejection AUC threshold', '\u2014', '30 mg\u00b7h/L', 'MPA TDM literature'),
    ]
    for row_data in pd_params:
        add_table_row(pd_table, row_data, font_size=8)

    # Table S-C: Allometric Exponents
    doc.add_paragraph('')
    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Table S-C. Allometric Scaling Exponents and Cross-References')
    run.bold = True
    run.font.size = Pt(10)

    allo_table = doc.add_table(rows=1, cols=4)
    allo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    allo_table.style = 'Light Grid Accent 1'
    for i, text in enumerate(['Parameter Type', 'Exponent', 'Equation', 'Reference']):
        cell = allo_table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    allo_data = [
        ('Distribution volumes (V\u1d9c, V\u209a, V\u2098\u209a\u2090\u1d4d)', '1.0', 'S8', cite(13)),
        ('Intercompartmental clearance (Q)', '0.75', 'S9', cite(13, 14)),
        ('Liver weight', '0.86', 'S10', cite(13)),
        ('Hepatic blood flow (Q\u2095)', '0.75', 'S11', cite(13)),
        ('Albumin\u2013binding', '1.2', 'S5', cite(7)),
        ('BSA (height component)', '0.725', 'S12', 'Standard BSA formula'),
        ('BSA (weight component)', '0.425', 'S12', 'Standard BSA formula'),
    ]
    for row_data in allo_data:
        add_table_row(allo_table, row_data, font_size=8)

    # ================================================================
    # NOTATION INDEX
    # ================================================================
    doc.add_page_break()
    doc.add_heading('Notation Index', level=1)

    doc.add_paragraph(
        'Complete listing of all mathematical symbols used in this supplement.'
    )

    notation_table = doc.add_table(rows=1, cols=3)
    notation_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    notation_table.style = 'Light Grid Accent 1'
    for i, text in enumerate(['Symbol', 'Definition', 'Units']):
        cell = notation_table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    notations = [
        ('CL\u2095', 'Hepatic clearance', 'L/h'),
        ('CL\u1d62\u2099\u209c', 'Total intrinsic clearance', 'L/h'),
        ('Q\u2095', 'Hepatic blood flow', 'L/h'),
        ('Q', 'Intercompartmental clearance', 'L/h'),
        ('f\u1d64', 'Unbound (free) fraction of MPA', 'dimensionless'),
        ('E', 'Hepatic extraction ratio', 'dimensionless'),
        ('V\u1d9c', 'Central compartment volume', 'L'),
        ('V\u209a', 'Peripheral compartment volume', 'L'),
        ('k\u2090', 'First-order absorption rate constant', '/h'),
        ('t\u2097\u2090\u1d4d', 'Absorption lag time', 'h'),
        ('F\u2092\u1d63\u2090\u2097', 'Oral bioavailability', 'dimensionless'),
        ('F\u1d4d\u2081', 'Gut wall bioavailability', 'dimensionless'),
        ('k\u2091\u1d62\u2097\u1d49', 'Biliary MPAG excretion rate', '/h'),
        ('k\u2091\u1d64\u209c', 'Gut bacterial deconjugation rate', '/h'),
        ('f\u1d31\u1d34\u1d30', 'Enterohepatic recirculation fraction', 'dimensionless'),
        ('f\u2092\u2090\u209c\u209a\u2081\u2091', 'OATP1B transporter function', 'dimensionless'),
        ('C\u2098\u209a\u2090', 'MPA plasma concentration (total)', 'mg/L'),
        ('C\u1da0\u1d63\u1d49\u1d49', 'Free MPA concentration', 'mg/L'),
        ('C\u2098\u209a\u2090\u1d4d', 'MPAG plasma concentration', 'mg/L'),
        ('AUC', 'Area under the concentration\u2013time curve', 'mg\u00b7h/L'),
        ('C\u2098\u2090\u2093', 'Maximum plasma concentration', 'mg/L'),
        ('t\u2098\u2090\u2093', 'Time to maximum concentration', 'h'),
        ('I', 'IMPDH-II inhibition fraction', 'dimensionless'),
        ('I\u2090\u1d65\u1d4d', 'Time-averaged IMPDH inhibition', 'dimensionless'),
        ('IC\u2085\u2080', 'Half-maximal inhibitory concentration', 'mg/L'),
        ('\u03b3', 'Hill coefficient (sigmoidicity)', 'dimensionless'),
        ('L\u209b\u209b', 'Steady-state lymphocyte count', '\u00d710\u2079/L'),
        ('P\u1d63\u1d49\u2c7c', 'Rejection probability', 'dimensionless'),
        ('P\u1d4d\u1d35', 'GI toxicity probability', 'dimensionless'),
        ('P\u2097\u1d49\u1d64\u2096\u2092', 'Leukopenia probability', 'dimensionless'),
        ('P\u1d62\u2099\u1da0', 'Infection probability', 'dimensionless'),
        ('P\u2090\u1d48\u1d65\u1d49\u1d63\u209b\u1d49', 'Composite adverse event probability', 'dimensionless'),
        ('TI', 'Therapeutic index', 'dimensionless'),
        ('GMFE', 'Geometric mean fold error', 'dimensionless'),
        ('FE', 'Individual fold error', 'dimensionless'),
        ('CV%', 'Coefficient of variation', '%'),
        ('BW', 'Body weight', 'kg'),
        ('LW', 'Liver weight', 'kg'),
        ('BSA', 'Body surface area', 'm\u00b2'),
        ('Alb', 'Serum albumin concentration', 'g/dL'),
        ('GFR', 'Glomerular filtration rate', 'mL/min'),
        ('\u03c4', 'Dosing interval', 'h'),
        ('D', 'Dose', 'mg'),
    ]
    for row_data in notations:
        add_table_row(notation_table, row_data, font_size=8)

    # Supplementary figures are maintained separately in Supplementary_figures.docx
    # (the dedicated, correctly-numbered figures file) and are intentionally NOT
    # embedded here to avoid duplication and numbering drift.

    # EndNote will populate the managed bibliography below this heading after
    # "Update Citations and Bibliography" is run in Word.
    doc.add_page_break()
    doc.add_heading('References', level=1)

    normalize_visible_notation(doc)

    # Save
    doc.save(OUTPUT_PATH)
    write_ris(REFERENCES, RIS_PATH)
    print(f"EndNote-ready supplement saved: {OUTPUT_PATH}")
    print(f"RIS reference file saved: {RIS_PATH}")
    print(f"Contains 33 numbered equations (S1-S33) across 10 sections")
    print(f"Includes 3 parameter tables (S-A, S-B, S-C) and notation index")


if __name__ == "__main__":
    create_supplement()
