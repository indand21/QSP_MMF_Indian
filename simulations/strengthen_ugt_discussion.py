# One-off editor: strengthen South Asian UGT pharmacogenetics discussion in the revision docx.
# Edits MPA_QSP_Manuscript_Revision.docx in place (backup: .preUGT.docx).
import os
from docx import Document
from docx.oxml.ns import qn

HERE = os.path.join(os.path.dirname(__file__), '..', 'outputs', 'manuscript_updated')
PATH = os.path.join(HERE, 'MPA_QSP_Manuscript_Revision.docx')

d = Document(PATH)

def set_paragraph_text(p, text):
    """Flatten paragraph to a single clean run, preserving paragraph properties."""
    p_el = p._p
    for child in list(p_el):
        if child.tag != qn('w:pPr'):
            p_el.remove(child)
    p.add_run(text)

# ---- new prose ----
PARA67 = (
    'Virtual populations of 500 patients each were generated to reflect contemporary transplant '
    'demographics under a matched albumin/CNI scenario. Both populations shared the following '
    'model-design assumptions: serum albumin 4.0 ± 0.4 g/dL (truncated normal, range 2.5–5.5 '
    'g/dL) and tacrolimus co-therapy at ~93%. These shared parameters were used to isolate body '
    'weight and UGT activity rather than to claim that albumin or CNI use is identical across all '
    'real-world cohorts. Differing parameters were Western body weight 78 ± 15 kg and Indian body '
    'weight 58 ± 12 kg; Western eGFR 55 ± 18 mL/min, Indian eGFR 50 ± 18 mL/min; Indian UGT1A9 '
    'activity 0.95 ± 0.18 (vs 1.0 ± 0.15), Indian UGT2B7 activity 0.92 ± 0.20 (vs 1.0 ± 0.15). '
    'The Indian weight distribution was chosen to represent a lower-weight/lean recipient scenario '
    'consistent with Indian and Asian transplant exposure reports [4-6]. UGT activities were sampled '
    'from log-normal distributions reflecting pharmacogenomic variability. The modest mean reductions '
    'modeled in Indian populations reflect an exploratory prior rather than an established effect-size '
    'estimate: direct genotype–activity data specific to South Asian MPA metabolism are limited, and '
    'reported UGT2B7 c.802C>T (rs7439366) T-allele frequencies in South Indian cohorts are '
    'approximately 37–40% [18, 24], below the ≈49–52% reported in European cohorts. The robustness '
    'of conclusions to this assumption is examined in the parameter-sensitivity analysis (Figure S1).'
)

PARA116 = (
    'A novel finding of this study is the rigorous quantification of the sources of greater '
    'interindividual variability in Indian patients (CV% 36.2% vs 30.3% for total AUC). The variance '
    'decomposition analysis (Figure S3) revealed that UGT1A9 enzyme activity is the dominant source '
    'of AUC variability in both populations (73% of explained variance in Indian, 68% in Western '
    'patients), followed by body weight (19% in both). The higher Indian variability is driven by the '
    'wider UGT1A9 activity distribution (SD 0.18 vs 0.15 on log-scale), an exploratory prior informed '
    'indirectly by South Asian UGT polymorphism data (e.g., UGT2B7 c.802C>T T-allele frequency '
    '≈37–40% in a South Indian non-transplant cohort [18]); direct South Asian-specific UGT1A9 '
    'functional-variant data for MPA remain sparse.'
)

PARA119 = (
    'These findings have direct clinical implications: the higher variability in Indian patients '
    'suggests a stronger rationale for routine therapeutic drug monitoring (TDM) and AUC-guided dose '
    'adjustment in this population [3, 5, 6, 17]. Furthermore, the pharmacogenetic architecture of MPA '
    'metabolism differs substantially between South Asian and Western populations in ways directly '
    'relevant to overexposure. The UGT1A9 promoter gain-of-function variants −275T>A and −2152C>T—'
    'which increase hepatic UGT1A9 expression, lower MPA exposure by ~20%, and have been associated '
    'with acute rejection in MMF/tacrolimus-treated Western recipients [23]—are essentially absent in '
    'South Asians (gnomAD minor-allele frequency ≈1% each, versus ≈5–6% in Europeans and 0% in East '
    'Asians [11, 22, 24]). The reduced-activity UGT1A9*3 (c.98T>C) allele, which conversely raises MPA '
    'exposure, is likewise near-absent in South Asians (<0.2% vs ≈1.4% in Europeans [24]). '
    'Consequently, the common European variants that shift MPA clearance are not operative at the '
    'population level in Indian patients, who lack the clearance-accelerating genotype that partially '
    'protects a subset (~10–15%) of Western recipients from overexposure. By contrast, the intronic '
    'UGT1A9 I399C>T variant—reported to influence MPA pharmacokinetics in Japanese renal transplant '
    'recipients [22]—reaches its highest global frequency in South Asians (≈33%, versus ≈30% in '
    'Europeans and ≈3% in East Asians [24]), although its functional impact on MPA exposure remains '
    'unresolved. Together with an intermediate UGT2B7 c.802C>T T-allele frequency (≈37–40% in South '
    'Indian cohorts [18], below the ≈49–52% reported in Europeans), these data indicate that the UGT '
    'variants most strongly linked to MPA exposure in Western and East Asian studies do not transfer '
    'directly to South Asians. Pharmacogenomic-guided dosing in Indian recipients will therefore '
    'require South Asian-specific functional-variant discovery rather than genotyping the established '
    'European markers, reinforcing AUC-guided TDM as the pragmatic near-term strategy [3, 17].'
)

# ---- locate & rewrite the three paragraphs by content signature ----
n67 = n116 = n119 = 0
for p in d.paragraphs:
    t = p.text
    if t.startswith('Virtual populations of 500 patients each') and 'UGT2B7 c.802C>T (rs7439366)' in t:
        set_paragraph_text(p, PARA67); n67 += 1
    elif t.startswith('A novel finding of this study') and 'UGT2B7 c.802C>T T-allele frequency' in t:
        set_paragraph_text(p, PARA116); n116 += 1
    elif t.startswith('These findings have direct clinical implications') and 'pharmacogenomic-guided dosing' in t:
        set_paragraph_text(p, PARA119); n119 += 1
assert (n67, n116, n119) == (1, 1, 1), (n67, n116, n119)

# ---- add bibliography entries 23 & 24 after entry 22 (Inoue) ----
REF23 = ('23.\tvan Schaik, R.H., et al., UGT1A9 -275T>A/-2152C>T polymorphisms correlate with low MPA '
         'exposure and acute rejection in MMF/tacrolimus-treated kidney transplant patients. Clinical '
         'Pharmacology and Therapeutics, 2009. 86(3): p. 319-327.')
REF24 = ('24.\tKarczewski, K.J., et al., The mutational constraint spectrum quantified from variation in '
         '141,456 humans. Nature, 2020. 581(7809): p. 434-443.')

inoue = None
for p in d.paragraphs:
    if p.text.strip().startswith('22.') and 'Inoue' in p.text:
        inoue = p
        break
assert inoue is not None, 'Inoue (ref 22) paragraph not found'
style = inoue.style

# insert after Inoue: build elements and place with addnext (reverse order to keep 23 then 24)
for txt in (REF24, REF23):
    new_p = inoue.insert_paragraph_before('')   # temp; we move it after inoue
    new_p.style = style
    new_p.add_run(txt)
    inoue._p.addnext(new_p._p)

d.save(PATH)
print('docx updated OK:', (n67, n116, n119))
