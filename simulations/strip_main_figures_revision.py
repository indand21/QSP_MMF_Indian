"""
Strip the four main figures from MPA_QSP_Manuscript_Revision.docx and relocate
their legends to a "Figure Legends" section at the very end of the document
(after the References bibliography).

The four main figures (Figure 1-4) are supplied to the journal as separate files,
so the embedded copies and their now-orphaned media parts are removed to shrink the
document. In-text "(Figure N)" references are intentionally left untouched.

This edits the docx IN PLACE (a .backup.docx is written first). It does NOT
regenerate the manuscript from the generator script.
"""

import os
import shutil

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn

project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG_DIR = os.path.join(project_root, "outputs", "manuscript_updated")
DOCX = os.path.join(FIG_DIR, "MPA_QSP_Manuscript_Revision.docx")
BACKUP = os.path.join(FIG_DIR, "MPA_QSP_Manuscript_Revision.backup.docx")


def find_image_paragraphs(doc):
    """Body paragraphs that contain an inline <w:drawing> (one per main figure)."""
    out = []
    for p in doc.paragraphs:
        if p._p.findall('.//' + qn('w:drawing')):
            out.append(p)
    return out


def find_legend_paragraphs(doc):
    """Return the four legend paragraphs in figure-number order (1->4)."""
    legends = []
    for n in ('1', '2', '3', '4'):
        prefix = 'Figure %s.' % n
        matches = [p for p in doc.paragraphs
                   if p.text.strip().startswith(prefix) and len(p.text.strip()) > 100]
        if len(matches) != 1:
            raise RuntimeError(
                "Expected exactly 1 legend paragraph for 'Figure %s.', found %d"
                % (n, len(matches)))
        legends.append(matches[0])
    return legends


def main():
    if not os.path.exists(DOCX):
        raise SystemExit("Not found: %s" % DOCX)

    shutil.copy2(DOCX, BACKUP)
    print("Backup written: %s" % BACKUP)

    doc = Document(DOCX)
    body = doc.element.body

    img_paras = find_image_paragraphs(doc)
    print("Image paragraphs found: %d" % len(img_paras))
    if len(img_paras) != 4:
        raise RuntimeError("Expected 4 image paragraphs, found %d" % len(img_paras))

    legend_paras = find_legend_paragraphs(doc)
    print("Legend paragraphs found: %d" % len(legend_paras))
    for p in legend_paras:
        print("  - %s..." % p.text.strip()[:40])

    # Detach legend elements first (keep references), then remove image paragraphs.
    legend_elems = [p._p for p in legend_paras]
    for el in legend_elems:
        el.getparent().remove(el)

    for p in img_paras:
        p._p.getparent().remove(p._p)

    # Append a "Figure Legends" heading at the logical end (before trailing sectPr),
    # then re-attach the four legend paragraphs in order.
    doc.add_heading('Figure Legends', level=1)

    sectPr = body.find(qn('w:sectPr'))
    if sectPr is None:
        raise RuntimeError("No trailing <w:sectPr> found in document body")
    for el in legend_elems:
        sectPr.addprevious(el)

    # Drop image relationships so the orphaned media parts are excluded on save.
    doc_part = doc.part
    image_rids = [rId for rId, rel in list(doc_part.rels.items())
                  if rel.reltype == RT.IMAGE]
    for rId in image_rids:
        doc_part.drop_rel(rId)
    print("Dropped %d image relationships: %s" % (len(image_rids), image_rids))

    doc.save(DOCX)
    print("Saved: %s" % DOCX)
    print("New size: %.2f MB" % (os.path.getsize(DOCX) / 1e6))


if __name__ == "__main__":
    main()
