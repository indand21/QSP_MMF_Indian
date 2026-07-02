"""
Generate Updated DOCX Manuscript: Contemporary Clinical Practice
===========================================================
Professional OMML equations, updated population design,
sensitivity analysis results, PD outcomes.
"""

import os
import numpy as np
from lxml import etree
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG_DIR = os.path.join(project_root, "outputs", "manuscript_updated")
OUTPUT_PATH = os.path.join(
    FIG_DIR,
    "MPA_QSP_Manuscript_Updated_albumin_CNI_fixed_endnote_ready.docx",
)
RIS_PATH = os.path.join(FIG_DIR, "MPA_QSP_References.ris")

MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
WORD_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


# ================================================================
# REFERENCES — validated against PubMed/DOI (see plan file)
# ================================================================

def build_references():
    """Return the manuscript references as structured dicts.

    Validated 2026-04-18 against PubMed. Corrections from the pre-validation
    list: Ref 8 (journal fixed to BJCP), Ref 9 (Pithukpakorn journal fixed
    to Pharmgenomics Pers Med), Ref 11 (journal fixed to JASN), Ref 13
    (Koloskoff journal fixed to Ther Drug Monit), Ref 16 (Kagaya replaced
    with Inoue 2007 which actually reports UGT1A9 genotypes).
    """
    refs = [
        dict(authors=['Staatz CE', 'Tett SE'],
             title='Clinical pharmacokinetics and pharmacodynamics of mycophenolate in solid organ transplant recipients',
             journal='Clinical Pharmacokinetics', journal_abbrev='Clin Pharmacokinet',
             year=2007, volume='46', issue='1', start_page='13', end_page='58',
             doi='10.2165/00003088-200746010-00002', pmid='17201457'),
        dict(authors=['de Winter BCM', 'van Gelder T', 'Sombogaard F', 'Shaw LM', 'van Hest RM', 'Mathot RAA'],
             title='Pharmacokinetic role of protein binding of mycophenolic acid and its glucuronide metabolite in renal transplant recipients',
             journal='Journal of Pharmacokinetics and Pharmacodynamics', journal_abbrev='J Pharmacokinet Pharmacodyn',
             year=2009, volume='36', issue='6', start_page='541', end_page='564',
             doi='10.1007/s10928-009-9136-6', pmid='19904584'),
        dict(authors=['Le Meur Y', 'B\u00fcchler M', 'Thierry A', 'Caillard S', 'Villemain F', 'Lavaud S',
                      'Etienne I', 'Westeel PF', 'Hurault de Ligny B', 'Rostaing L', 'Thervet E',
                      'Szelag JC', 'Rerolle JP', 'Rousseau A', 'Touchard G', 'Marquet P'],
             title='Individualized mycophenolate mofetil dosing based on drug exposure significantly improves patient outcomes after renal transplantation',
             journal='American Journal of Transplantation', journal_abbrev='Am J Transplant',
             year=2007, volume='7', issue='11', start_page='2496', end_page='2503',
             doi='10.1111/j.1600-6143.2007.01983.x', pmid='17908276'),
        dict(authors=['Glander P', 'Hambach P', 'Braun KP', 'Fritsche L', 'Giessing M', 'Mai I',
                      'Einecke G', 'Waiser J', 'Neumayer HH', 'Budde K'],
             title='Pre-transplant inosine monophosphate dehydrogenase activity is associated with clinical outcome after renal transplantation',
             journal='American Journal of Transplantation', journal_abbrev='Am J Transplant',
             year=2004, volume='4', issue='12', start_page='2045', end_page='2051',
             doi='10.1111/j.1600-6143.2004.00617.x', pmid='15575908'),
        dict(authors=['van Gelder T', 'Le Meur Y', 'Shaw LM', 'Oellerich M', 'DeNofrio D', 'Holt C',
                      'Holt DW', 'Kaplan B', 'Kuypers D', 'Meiser B', 'Toenshoff B', 'Mamelok RD'],
             title='Therapeutic drug monitoring of mycophenolate mofetil in transplantation',
             journal='Therapeutic Drug Monitoring', journal_abbrev='Ther Drug Monit',
             year=2006, volume='28', issue='2', start_page='145', end_page='154',
             doi='10.1097/01.ftd.0000199358.80013.bd', pmid='16628123'),
        dict(authors=['Colom H', 'Lloberas N', 'Andreu F', 'Caldes A', 'Torras J', 'Oppenheimer F',
                      'Sanchez-Plumed J', 'Gentil MA', 'Kuypers DR', 'Brunet M', 'Ekberg H', 'Grinyo JM'],
             title='Pharmacokinetic modeling of enterohepatic circulation of mycophenolic acid in renal transplant recipients',
             journal='Kidney International', journal_abbrev='Kidney Int',
             year=2014, volume='85', issue='6', start_page='1434', end_page='1443',
             doi='10.1038/ki.2013.517', pmid='24402086'),
        dict(authors=['Yau WP', 'Vathsala A', 'Lou HX', 'Chan E'],
             title='Is a standard fixed dose of mycophenolate mofetil ideal for all patients?',
             journal='Nephrology Dialysis Transplantation', journal_abbrev='Nephrol Dial Transplant',
             year=2007, volume='22', issue='12', start_page='3638', end_page='3645',
             doi='10.1093/ndt/gfm468', pmid='17640939'),
        dict(authors=['Yu ZC', 'Zhou PJ', 'Xu D', 'Wang XH', 'Chen HZ'],
             title='Investigation on pharmacokinetics of mycophenolic acid in Chinese adult renal transplant patients',
             journal='British Journal of Clinical Pharmacology', journal_abbrev='Br J Clin Pharmacol',
             year=2006, volume='62', issue='4', start_page='446', end_page='452',
             doi='10.1111/j.1365-2125.2006.02626.x', pmid='16995865'),
        dict(authors=['Pithukpakorn M', 'Tiwawanwong T', 'Lalerd Y', 'Assawamakin A',
                      'Premasathian N', 'Tasanarong A', 'Thongnoppakhun W', 'Vongwiwatana A'],
             title='Mycophenolic acid AUC in Thai kidney transplant recipients receiving low dose mycophenolate and its association with UGT2B7 polymorphisms',
             journal='Pharmacogenomics and Personalized Medicine', journal_abbrev='Pharmgenomics Pers Med',
             year=2014, volume='7', issue='', start_page='379', end_page='385',
             doi='10.2147/PGPM.S72760', pmid='25540593'),
        dict(authors=['Shaw LM', 'Korecka M', 'Venkataramanan R', 'Goldberg L', 'Bloom R', 'Brayman KL'],
             title='Mycophenolic acid pharmacodynamics and pharmacokinetics provide a basis for rational monitoring strategies',
             journal='American Journal of Transplantation', journal_abbrev='Am J Transplant',
             year=2003, volume='3', issue='5', start_page='534', end_page='542',
             doi='10.1034/j.1600-6143.2003.00079.x', pmid='12752309'),
        dict(authors=['van Hest RM', 'Mathot RAA', 'Pescovitz MD', 'Gordon R', 'Mamelok RD', 'van Gelder T'],
             title='Explaining variability in mycophenolic acid exposure to optimize mycophenolate mofetil dosing: a population pharmacokinetic meta-analysis of mycophenolic acid in renal transplant recipients',
             journal='Journal of the American Society of Nephrology', journal_abbrev='J Am Soc Nephrol',
             year=2006, volume='17', issue='3', start_page='871', end_page='880',
             doi='10.1681/ASN.2005101070', pmid='16452491'),
        dict(authors=['Sombogaard F', 'van Schaik RHN', 'Mathot RAA', 'Budde K', 'van der Werf M',
                      'Vulto AG', 'Weimar W', 'Glander P', 'Essioux L', 'van Gelder T'],
             title='Interpatient variability in IMPDH activity in MMF-treated renal transplant patients is correlated with IMPDH type II 3757T>C polymorphism',
             journal='Pharmacogenetics and Genomics', journal_abbrev='Pharmacogenet Genomics',
             year=2009, volume='19', issue='8', start_page='626', end_page='634',
             doi='10.1097/FPC.0b013e32832f5f1b', pmid='19617864'),
        dict(authors=['Koloskoff K', 'Panwar R', 'Rathi M', 'Mathew S', 'Sharma A', 'Marquet P',
                      'Benito S', 'Woillard JB', 'Pattanaik S'],
             title='Population pharmacokinetics and limited sampling strategy of mycophenolate mofetil for Indian patients with lupus nephritis',
             journal='Therapeutic Drug Monitoring', journal_abbrev='Ther Drug Monit',
             year=2024, volume='46', issue='5', start_page='567', end_page='574',
             doi='10.1097/FTD.0000000000001213', pmid='38723153'),
        dict(authors=['Jiao Z', 'Ding JJ', 'Shen J', 'Liang HQ', 'Zhong LJ', 'Wang Y', 'Zhong MK', 'Lu WY'],
             title='Population pharmacokinetic modelling for enterohepatic circulation of mycophenolic acid in healthy Chinese and the influence of polymorphisms in UGT1A9',
             journal='British Journal of Clinical Pharmacology', journal_abbrev='Br J Clin Pharmacol',
             year=2008, volume='65', issue='6', start_page='893', end_page='907',
             doi='10.1111/j.1365-2125.2008.03109.x', pmid='18279479'),
        dict(authors=['Inoue K', 'Miura M', 'Satoh S', 'Kagaya H', 'Saito M', 'Habuchi T', 'Suzuki T'],
             title='Influence of UGT1A7 and UGT1A9 intronic I399 genetic polymorphisms on mycophenolic acid pharmacokinetics in Japanese renal transplant recipients',
             journal='Therapeutic Drug Monitoring', journal_abbrev='Ther Drug Monit',
             year=2007, volume='29', issue='3', start_page='299', end_page='304',
             doi='10.1097/FTD.0b013e3180686146', pmid='17529886'),
        dict(authors=['Anderson BJ', 'Holford NHG'],
             title='Mechanism-based concepts of size and maturity in pharmacokinetics',
             journal='Annual Review of Pharmacology and Toxicology', journal_abbrev='Annu Rev Pharmacol Toxicol',
             year=2008, volume='48', issue='', start_page='303', end_page='332',
             doi='10.1146/annurev.pharmtox.48.113006.094708', pmid='17914927'),
        dict(authors=['West GB', 'Brown JH', 'Enquist BJ'],
             title='A general model for the origin of allometric scaling laws in biology',
             journal='Science', journal_abbrev='Science',
             year=1997, volume='276', issue='5309', start_page='122', end_page='126',
             doi='10.1126/science.276.5309.122', pmid='9082983'),
        dict(authors=['Picard N', 'Ratanasavanh D', 'Pr\u00e9maud A', 'Le Meur Y', 'Marquet P'],
             title='Identification of the UDP-glucuronosyltransferase isoforms involved in mycophenolic acid phase II metabolism',
             journal='Drug Metabolism and Disposition', journal_abbrev='Drug Metab Dispos',
             year=2005, volume='33', issue='1', start_page='139', end_page='146',
             doi='10.1124/dmd.104.001651', pmid='15470161'),
        dict(authors=['Venkatraman S', 'Ramasamy K', 'Nair PP'],
             title='Genetic polymorphisms of microsomal epoxide hydrolase and UDP-glucuronosyltransferase (UGT) and its effects on plasma carbamazepine levels and metabolic ratio in persons with epilepsy of South India: a cross-sectional genetic association study',
             journal='Indian Journal of Pharmacology', journal_abbrev='Indian J Pharmacol',
             year=2023, volume='55', issue='3', start_page='149', end_page='154',
             doi='10.4103/ijp.ijp_228_22', pmid='37555408'),
        dict(authors=['Allison AC', 'Eugui EM'],
             title='Mycophenolate mofetil and its mechanisms of action',
             journal='Immunopharmacology', journal_abbrev='Immunopharmacology',
             year=2000, volume='47', issue='2-3', start_page='85', end_page='118',
             doi='10.1016/S0162-3109(00)00188-0', pmid='10878285'),
        dict(authors=['Singh S', 'Panwar R', 'Naithani P', 'Kathik NVKN', 'Sharma N',
                      'Gowda N', 'Kenwar DB', 'Kumar SPS', 'Singh S', 'Pattanaik S',
                      'Sharma A'],
             title='Exposure to mycophenolic acid at standard prescribed doses in renal transplantation recipients and clinical outcomes in the early posttransplantation period',
             journal='Indian Journal of Urology', journal_abbrev='Indian J Urol',
             year=2025, volume='41', issue='4', start_page='287', end_page='295',
             doi='10.4103/iju.iju_43_25', pmid='41112717'),
        dict(authors=['Sarangi SC', 'Reeta KH', 'Agarwal SK', 'Kaleekal T', 'Guleria S',
                      'Gupta YK'],
             title='A pilot study on area under curve of mycophenolic acid as a guide for its optimal use in renal transplant recipients',
             journal='Indian Journal of Medical Research', journal_abbrev='Indian J Med Res',
             year=2012, volume='135', issue='1', start_page='84', end_page='91',
             doi='10.4103/0971-5916.93429', pmid='22382188'),
        dict(authors=['van Schaik RHN', 'van Agteren M', 'de Fijter JW', 'Hartmann A', 'Schmidt J',
                      'Budde K', 'Kuypers D', 'Le Meur Y', 'van der Werf M', 'Mamelok R',
                      'van Gelder T'],
             title='UGT1A9 -275T>A/-2152C>T polymorphisms correlate with low MPA exposure and acute rejection in MMF/tacrolimus-treated kidney transplant patients',
             journal='Clinical Pharmacology and Therapeutics', journal_abbrev='Clin Pharmacol Ther',
             year=2009, volume='86', issue='3', start_page='319', end_page='327',
             doi='10.1038/clpt.2009.83', pmid='19494809'),
        dict(authors=['Karczewski KJ', 'Francioli LC', 'Tiao G', 'Cummings BB', 'Alföldi J',
                      'Wang Q', 'Collins RL', 'Laricchia KM', 'Ganna A', 'Birnbaum DP', 'MacArthur DG'],
             title='The mutational constraint spectrum quantified from variation in 141,456 humans',
             journal='Nature', journal_abbrev='Nature',
             year=2020, volume='581', issue='7809', start_page='434', end_page='443',
             doi='10.1038/s41586-020-2308-7', pmid='32461654'),
    ]
    final_bibliography_order = [
        1,   # Staatz
        20,  # Allison
        5,   # van Gelder
        7,   # Yau
        21,  # Singh
        22,  # Sarangi
        2,   # de Winter
        6,   # Colom
        9,   # Pithukpakorn
        11,  # van Hest
        14,  # Jiao
        18,  # Picard
        16,  # Anderson
        17,  # West
        4,   # Glander
        12,  # Sombogaard
        3,   # Le Meur
        19,  # Venkatraman
        24,  # Karczewski (gnomAD)
        10,  # Shaw
        8,   # Yu
        13,  # Koloskoff
        23,  # van Schaik
        15,  # Inoue
    ]
    return [refs[i - 1] for i in final_bibliography_order]


def format_vancouver(ref, author_limit=6):
    """Format a reference dict as a Vancouver-style plain-text citation."""
    authors = ref['authors']
    if len(authors) > author_limit:
        author_str = ', '.join(authors[:author_limit]) + ', et al'
    else:
        author_str = ', '.join(authors)
    title = ref['title'].rstrip('.')
    parts = [f'{author_str}. {title}. {ref["journal_abbrev"]}. {ref["year"]}']
    loc = ''
    if ref.get('volume'):
        loc = f';{ref["volume"]}'
        if ref.get('issue'):
            loc += f'({ref["issue"]})'
        if ref.get('start_page'):
            end = ref.get('end_page') or ref['start_page']
            loc += f':{ref["start_page"]}\u2013{end}'
    parts.append(loc + '.')
    return ''.join(parts)


def _first_author_surname(authors):
    """Extract surname from 'LastName FI' — keeps particles like 'de', 'van', 'Le'."""
    tokens = authors[0].split()
    if len(tokens) > 1 and tokens[-1].isupper() and len(tokens[-1]) <= 4:
        return ' '.join(tokens[:-1])
    return authors[0]


def _format_ris_author(author):
    """Return RIS-friendly 'Surname, Initials' author text."""
    tokens = author.split()
    if len(tokens) > 1 and tokens[-1].isupper() and len(tokens[-1]) <= 4:
        return f'{" ".join(tokens[:-1])}, {tokens[-1]}'
    return author


def _endnote_temp_marker(ref_nums, references):
    """Build an EndNote temporary citation marker for a list of ref numbers.

    EndNote format: {FirstAuthor, Year, unique text; FirstAuthor, Year, unique text}
    When the user runs EndNote's 'Update Citations and Bibliography',
    each temporary citation is resolved against their EndNote library. The
    unique text is DOI when available, otherwise PMID or title. This avoids
    record numbers, which are specific to a single EndNote library.
    """
    parts = []
    for n in ref_nums:
        ref = references[n - 1]
        surname = _first_author_surname(ref['authors'])
        unique_text = ref.get('doi') or ref.get('pmid') or ref['title']
        parts.append(f'{surname}, {ref["year"]}, {unique_text}')
    return '{' + '; '.join(parts) + '}'


def transform_bracket_citations_in_doc(doc, references):
    """Replace every [n,m,...] bracket citation in the document with an
    EndNote temporary citation marker. Operates on plain-text runs only,
    leaving equations (OMML), field codes, and reference-list entries untouched."""
    import re
    bracket_re = re.compile(r'\[([\d,\u2013\-\s]+)\]')

    def repl(match):
        inner = match.group(1).replace('\u2013', '-').replace(' ', '')
        nums = []
        for chunk in inner.split(','):
            if '-' in chunk:
                lo, hi = chunk.split('-', 1)
                nums.extend(range(int(lo), int(hi) + 1))
            elif chunk:
                nums.append(int(chunk))
        if not nums:
            return match.group(0)
        return _endnote_temp_marker(nums, references)

    def transform_paragraph(para):
        for run in para.runs:
            if '[' in run.text and ']' in run.text:
                new_text = bracket_re.sub(repl, run.text)
                if new_text != run.text:
                    run.text = new_text

    for para in doc.paragraphs:
        transform_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    transform_paragraph(para)


def remap_numeric_citations_in_doc(doc, old_to_new):
    """Remap bracketed numeric citations after manual reference reordering."""
    import re

    bracket_re = re.compile(r'\[([\d,\u2013\-\s]+)\]')

    def expand(inner):
        nums = []
        for chunk in inner.replace('\u2013', '-').replace(' ', '').split(','):
            if not chunk:
                continue
            if '-' in chunk:
                lo, hi = chunk.split('-', 1)
                nums.extend(range(int(lo), int(hi) + 1))
            else:
                nums.append(int(chunk))
        return nums

    def compress(nums):
        vals = sorted({old_to_new.get(n, n) for n in nums})
        ranges = []
        i = 0
        while i < len(vals):
            start = vals[i]
            end = start
            while i + 1 < len(vals) and vals[i + 1] == end + 1:
                i += 1
                end = vals[i]
            ranges.append(str(start) if start == end else f'{start}\u2013{end}')
            i += 1
        return ','.join(ranges)

    def repl(match):
        nums = expand(match.group(1))
        return f'[{compress(nums)}]' if nums else match.group(0)

    def transform_paragraph(para):
        for run in para.runs:
            if '[' in run.text and ']' in run.text:
                new_text = bracket_re.sub(repl, run.text)
                if new_text != run.text:
                    run.text = new_text

    for para in doc.paragraphs:
        transform_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    transform_paragraph(para)


def write_ris(references, out_path):
    """Write an EndNote/Reference Manager compatible RIS file.

    Import in EndNote: File \u2192 Import \u2192 File \u2192 Import Option:
    Reference Manager (RIS). Fields emitted: TY, AU, TI, T2, JA, PY, VL, IS,
    SP, EP, DO, AN (PMID), UR, N1 (reference number), ER.
    """
    lines = []
    for i, ref in enumerate(references, 1):
        lines.append('TY  - JOUR')
        for author in ref['authors']:
            lines.append(f'AU  - {_format_ris_author(author)}')
        lines.append(f'TI  - {ref["title"]}')
        lines.append(f'T2  - {ref["journal"]}')
        lines.append(f'JA  - {ref["journal_abbrev"]}')
        lines.append(f'PY  - {ref["year"]}')
        if ref.get('volume'):
            lines.append(f'VL  - {ref["volume"]}')
        if ref.get('issue'):
            lines.append(f'IS  - {ref["issue"]}')
        if ref.get('start_page'):
            lines.append(f'SP  - {ref["start_page"]}')
        if ref.get('end_page'):
            lines.append(f'EP  - {ref["end_page"]}')
        if ref.get('doi'):
            lines.append(f'DO  - {ref["doi"]}')
            lines.append(f'UR  - https://doi.org/{ref["doi"]}')
        if ref.get('pmid'):
            lines.append(f'AN  - {ref["pmid"]}')
        lines.append(f'N1  - Reference #{i} in MPA QSP manuscript')
        lines.append('ER  - ')
        lines.append('')
    with open(out_path, 'w', encoding='utf-8', newline='\r\n') as f:
        f.write('\n'.join(lines))


# ================================================================
# OMML EQUATION HELPERS
# ================================================================

def math_sub(parent, tag):
    return etree.SubElement(parent, f'{{{MATH_NS}}}{tag}')

def word_sub(parent, tag):
    return etree.SubElement(parent, f'{{{WORD_NS}}}{tag}')

def math_run(parent, text, italic=True, bold=False, font_name='Cambria Math'):
    r = math_sub(parent, 'r')
    if not italic or bold:
        rPr = math_sub(r, 'rPr')
        if not italic:
            sty = math_sub(rPr, 'sty')
            sty.set(f'{{{MATH_NS}}}val', 'p')
    w_rPr = word_sub(r, 'rPr')
    rFonts = word_sub(w_rPr, 'rFonts')
    rFonts.set(f'{{{WORD_NS}}}ascii', font_name)
    rFonts.set(f'{{{WORD_NS}}}hAnsi', font_name)
    if bold:
        word_sub(w_rPr, 'b')
    t = math_sub(r, 't')
    t.text = text
    return r

def math_fraction(parent):
    f = math_sub(parent, 'f')
    fPr = math_sub(f, 'fPr')
    typ = math_sub(fPr, 'type')
    typ.set(f'{{{MATH_NS}}}val', 'bar')
    num = math_sub(f, 'num')
    den = math_sub(f, 'den')
    return f, num, den

def math_subscript(parent, base_text, sub_text, base_italic=True, sub_italic=True):
    sSub = math_sub(parent, 'sSub')
    math_sub(sSub, 'sSubPr')
    e = math_sub(sSub, 'e')
    math_run(e, base_text, italic=base_italic)
    sub = math_sub(sSub, 'sub')
    math_run(sub, sub_text, italic=sub_italic)
    return sSub

def math_superscript(parent, base_text, sup_text, base_italic=True):
    sSup = math_sub(parent, 'sSup')
    math_sub(sSup, 'sSupPr')
    e = math_sub(sSup, 'e')
    math_run(e, base_text, italic=base_italic)
    sup = math_sub(sSup, 'sup')
    math_run(sup, sup_text)
    return sSup

def math_subsup(parent, base_text, sub_text, sup_text):
    sSubSup = math_sub(parent, 'sSubSup')
    math_sub(sSubSup, 'sSubSupPr')
    e = math_sub(sSubSup, 'e')
    math_run(e, base_text)
    sub_elem = math_sub(sSubSup, 'sub')
    math_run(sub_elem, sub_text)
    sup_elem = math_sub(sSubSup, 'sup')
    math_run(sup_elem, sup_text)
    return sSubSup

def math_delimiters(parent, open_char='(', close_char=')'):
    d = math_sub(parent, 'd')
    dPr = math_sub(d, 'dPr')
    begChr = math_sub(dPr, 'begChr')
    begChr.set(f'{{{MATH_NS}}}val', open_char)
    endChr = math_sub(dPr, 'endChr')
    endChr.set(f'{{{MATH_NS}}}val', close_char)
    e = math_sub(d, 'e')
    return e

def create_display_equation(doc, builder_func, number=None):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    oMathPara = etree.SubElement(para._element, f'{{{MATH_NS}}}oMathPara')
    oMath = math_sub(oMathPara, 'oMath')
    builder_func(oMath)
    if number:
        run = para.add_run(f'\t({number})')
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    return para

def add_inline_math(paragraph, builder_func):
    oMath = etree.SubElement(paragraph._element, f'{{{MATH_NS}}}oMath')
    builder_func(oMath)
    return oMath


# ================================================================
# EQUATION BUILDERS
# ================================================================

def eq_well_stirred(oMath):
    math_subscript(oMath, 'CL', 'h', base_italic=True, sub_italic=False)
    math_run(oMath, ' = ', italic=False)
    f, num, den = math_fraction(oMath)
    math_subscript(num, 'Q', 'h')
    math_run(num, ' \u00b7 ')
    math_subscript(num, 'f', 'u')
    math_run(num, ' \u00b7 ')
    math_subscript(num, 'CL', 'int')
    math_subscript(den, 'Q', 'h')
    math_run(den, ' + ')
    math_subscript(den, 'f', 'u')
    math_run(den, ' \u00b7 ')
    math_subscript(den, 'CL', 'int')

def eq_extraction_ratio(oMath):
    math_run(oMath, 'E')
    math_run(oMath, ' = ', italic=False)
    f, num, den = math_fraction(oMath)
    math_subscript(num, 'f', 'u')
    math_run(num, ' \u00b7 ')
    math_subscript(num, 'CL', 'int')
    math_subscript(den, 'Q', 'h')
    math_run(den, ' + ')
    math_subscript(den, 'f', 'u')
    math_run(den, ' \u00b7 ')
    math_subscript(den, 'CL', 'int')
    math_run(oMath, ' \u2248 0.13', italic=False)

def eq_clh_approx(oMath):
    math_run(oMath, 'For low ', italic=False)
    math_run(oMath, 'E')
    math_run(oMath, ':  ', italic=False)
    math_subscript(oMath, 'CL', 'h')
    math_run(oMath, ' \u2248 ', italic=False)
    math_subscript(oMath, 'f', 'u')
    math_run(oMath, ' \u00b7 ')
    math_subscript(oMath, 'CL', 'int')

def eq_free_fraction(oMath):
    math_subscript(oMath, 'f', 'u')
    math_run(oMath, ' = ', italic=False)
    sSub = math_sub(oMath, 'sSub')
    math_sub(sSub, 'sSubPr')
    e = math_sub(sSub, 'e')
    math_run(e, 'f')
    sub = math_sub(sSub, 'sub')
    math_run(sub, 'u,ref', italic=False)
    math_run(oMath, ' \u00b7 ', italic=False)
    sSup = math_sub(oMath, 'sSup')
    math_sub(sSup, 'sSupPr')
    e = math_sub(sSup, 'e')
    inner = math_delimiters(e)
    f, num, den = math_fraction(inner)
    math_subscript(num, 'Alb', 'ref')
    math_run(den, 'Alb')
    sup = math_sub(sSup, 'sup')
    math_run(sup, '1.2', italic=False)

def eq_allometric_volume(oMath):
    math_subscript(oMath, 'V', 'i')
    math_run(oMath, ' = ', italic=False)
    math_subscript(oMath, 'V', 'ref')
    math_run(oMath, ' \u00b7 ', italic=False)
    sSup = math_sub(oMath, 'sSup')
    math_sub(sSup, 'sSupPr')
    e = math_sub(sSup, 'e')
    inner = math_delimiters(e)
    f, num, den = math_fraction(inner)
    math_subscript(num, 'BW', 'i')
    math_subscript(den, 'BW', 'ref')
    sup = math_sub(sSup, 'sup')
    math_run(sup, '1.0', italic=False)

def eq_allometric_clearance(oMath):
    math_subscript(oMath, 'CL', 'i')
    math_run(oMath, ' = ', italic=False)
    math_subscript(oMath, 'CL', 'ref')
    math_run(oMath, ' \u00b7 ', italic=False)
    sSup = math_sub(oMath, 'sSup')
    math_sub(sSup, 'sSupPr')
    e = math_sub(sSup, 'e')
    inner = math_delimiters(e)
    f, num, den = math_fraction(inner)
    math_subscript(num, 'BW', 'i')
    math_subscript(den, 'BW', 'ref')
    sup = math_sub(sSup, 'sup')
    math_run(sup, '0.75', italic=False)

def eq_emax_impdh(oMath):
    math_run(oMath, 'I')
    math_run(oMath, ' = ', italic=False)
    f, num, den = math_fraction(oMath)
    math_subscript(num, 'E', 'max')
    math_run(num, ' \u00b7 ')
    sSubSup = math_sub(num, 'sSubSup')
    math_sub(sSubSup, 'sSubSupPr')
    e = math_sub(sSubSup, 'e')
    math_run(e, 'C')
    sub_e = math_sub(sSubSup, 'sub')
    math_run(sub_e, 'free', italic=False)
    sup_e = math_sub(sSubSup, 'sup')
    math_run(sup_e, '\u03b3')
    sSubSup2 = math_sub(den, 'sSubSup')
    math_sub(sSubSup2, 'sSubSupPr')
    e2 = math_sub(sSubSup2, 'e')
    math_run(e2, 'IC', italic=False)
    sub_e2 = math_sub(sSubSup2, 'sub')
    math_run(sub_e2, '50', italic=False)
    sup_e2 = math_sub(sSubSup2, 'sup')
    math_run(sup_e2, '\u03b3')
    math_run(den, ' + ')
    sSubSup3 = math_sub(den, 'sSubSup')
    math_sub(sSubSup3, 'sSubSupPr')
    e3 = math_sub(sSubSup3, 'e')
    math_run(e3, 'C')
    sub_e3 = math_sub(sSubSup3, 'sub')
    math_run(sub_e3, 'free', italic=False)
    sup_e3 = math_sub(sSubSup3, 'sup')
    math_run(sup_e3, '\u03b3')

def eq_therapeutic_index(oMath):
    math_run(oMath, 'TI', italic=False)
    math_run(oMath, ' = ', italic=False)
    sSub = math_sub(oMath, 'sSub')
    math_sub(sSub, 'sSubPr')
    e = math_sub(sSub, 'e')
    math_run(e, 'IMPDH', italic=False)
    sub = math_sub(sSub, 'sub')
    math_run(sub, 'avg', italic=False)
    math_run(oMath, ' \u00d7 ', italic=False)
    inner = math_delimiters(oMath)
    math_run(inner, '1 \u2212 ', italic=False)
    math_subscript(inner, 'P', 'adverse')

def eq_ode_central(oMath):
    f_lhs, num_lhs, den_lhs = math_fraction(oMath)
    math_run(num_lhs, 'dC')
    math_run(den_lhs, 'dt')
    math_run(oMath, ' = ', italic=False)
    f_rhs, num_rhs, den_rhs = math_fraction(oMath)
    math_subscript(num_rhs, 'R', 'abs')
    math_run(num_rhs, ' + ')
    math_subscript(num_rhs, 'R', 'dist,in')
    math_run(num_rhs, ' + ')
    math_subscript(num_rhs, 'R', 'EHC')
    math_run(num_rhs, ' \u2212 ')
    math_subscript(num_rhs, 'R', 'dist,out')
    math_run(num_rhs, ' \u2212 ')
    math_subscript(num_rhs, 'R', 'met')
    math_subscript(den_rhs, 'V', 'c')

def eq_auc_target(oMath):
    math_run(oMath, '30', italic=False)
    math_run(oMath, ' \u2264 ', italic=False)
    math_subscript(oMath, 'AUC', '0\u201312h')
    math_run(oMath, ' \u2264 ', italic=False)
    math_run(oMath, '60  mg\u00b7h/L', italic=False)

def eq_gmfe(oMath):
    math_run(oMath, 'GMFE', italic=False)
    math_run(oMath, ' = exp', italic=False)
    inner = math_delimiters(oMath)
    math_run(inner, 'mean', italic=False)
    inner2 = math_delimiters(inner)
    inner3 = math_delimiters(inner2, '|', '|')
    math_run(inner3, 'log', italic=False)
    inner4 = math_delimiters(inner3)
    f, num, den = math_fraction(inner4)
    math_subscript(num, 'AUC', 'pred')
    math_subscript(den, 'AUC', 'obs')

def eq_cv(oMath):
    math_run(oMath, 'CV%', italic=False)
    math_run(oMath, ' = ', italic=False)
    f, num, den = math_fraction(oMath)
    math_run(num, 'SD')
    math_run(den, 'Mean')
    math_run(oMath, ' \u00d7 100', italic=False)


# ================================================================
# TABLE HELPERS
# ================================================================

def add_table_row(table, cells_data, bold=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        para = cell.paragraphs[0]
        run = para.add_run(str(text))
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        if bold:
            run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
    return row


def add_figure_legend(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    run.italic = True
    return p


# Supplementary figures are numbered (S1-S5) in order of first citation in the
# manuscript body. The PNG filenames retain their original content-based names.
SUPPLEMENTARY_FIGURES = [
    {
        "number": "S1",
        "filename": "FigureS5_ParameterSensitivity.png",
        "title": "Parameter uncertainty sensitivity analysis",
        "legend": (
            "Figure S1. Parameter uncertainty sensitivity analysis. "
            "One-at-a-time local sensitivity analysis of total MPA AUC_0-12h was performed at the "
            "Indian mean-patient baseline receiving 1 g MMF BID. Six structural PBPK parameters were "
            "perturbed by +/-20% from their nominal values: reference free fraction at albumin 4.0 g/dL, "
            "UGT1A9 intrinsic clearance, UGT2B7 intrinsic clearance, central volume, absorption rate, "
            "and biliary efflux rate. Bars show the percentage change in total AUC relative to the "
            "baseline simulation, with parameters ordered by effect size. The analysis identifies "
            "the reference free fraction and UGT1A9 intrinsic clearance as the dominant local "
            "uncertainty drivers while showing that single-parameter perturbations remain smaller "
            "than the population-level Indian-vs-Western exposure difference."
        ),
    },
    {
        "number": "S2",
        "filename": "FigureS1_Sensitivity.png",
        "title": "Factor decomposition sensitivity analysis",
        "legend": (
            "Figure S2. Factor decomposition sensitivity analysis. "
            "This figure decomposes the simulated exposure difference between the Western and Indian "
            "virtual transplant cohorts under the matched contemporary-practice scenario. Panel A "
            "shows the total MPA AUC_0-12h waterfall from the Western baseline to the Indian total, "
            "with the increment attributable to the combined body-weight, UGT-activity, and eGFR "
            "assumptions. Panel B shows the corresponding decomposition for free MPA AUC_0-12h. "
            "Panel C compares the body-weight distributions used to construct the Western and Indian "
            "virtual populations. Albumin and CNI co-therapy were deliberately matched between "
            "populations in this scenario, so they are not represented as between-population "
            "drivers in the waterfall."
        ),
    },
    {
        "number": "S3",
        "filename": "FigureS4_Variability_Decomposition.png",
        "title": "Variability decomposition analysis",
        "legend": (
            "Figure S3. Variability decomposition analysis. "
            "This figure quantifies sources of interindividual variability in simulated total MPA "
            "AUC_0-12h for Western and Indian virtual populations. Panel A shows univariate Pearson "
            "correlations between AUC and model covariates. Panel B shows partial correlations after "
            "adjustment for the other covariates. Panel C ranks Indian-population variance attribution "
            "by the reduction in AUC CV% obtained when each covariate is fixed at its population mean. "
            "Panel D compares the same variance-attribution metric between Western and Indian "
            "populations. Panel E overlays simulated weight-AUC relationships with the theoretical "
            "AUC proportional to BW^-0.75 relationship, and Panel F shows the derivative of that "
            "relationship, demonstrating the steeper exposure sensitivity at the Indian mean body "
            "weight compared with the Western mean body weight."
        ),
    },
    {
        "number": "S4",
        "filename": "FigureS2_Profiles.png",
        "title": "Representative PK/PD profiles",
        "legend": (
            "Figure S4. Representative PK/PD profiles. "
            "Representative steady-state 12-hour dosing-interval simulations are shown for a Western "
            "reference patient receiving 1000 mg MMF BID, an Indian reference patient receiving "
            "1000 mg MMF BID, the same Indian reference patient receiving 750 mg BID, and the same "
            "Indian reference patient receiving the weight-based 12 mg/kg BID strategy rounded to an "
            "available tablet strength. Panel A shows total MPA concentration-time profiles, Panel B "
            "shows free MPA concentration-time profiles, and Panel C shows the corresponding IMPDH-II "
            "inhibition profiles. Panel D summarizes model-derived clinical outcome metrics, including "
            "rejection probability, GI toxicity probability, and the therapeutic index, illustrating "
            "how weight-based dose reduction lowers toxicity while preserving pharmacodynamic activity."
        ),
    },
    {
        "number": "S5",
        "filename": "FigureS3_PD_Heatmaps.png",
        "title": "PD outcome heatmaps by weight and AUC",
        "legend": (
            "Figure S5. PD outcome heatmaps by body weight and total MPA exposure. "
            "Each point represents a virtual patient plotted by body weight and total MPA AUC_0-12h, "
            "with color indicating the predicted pharmacodynamic outcome. The three panels show GI "
            "toxicity probability, leukopenia probability, and therapeutic index, respectively. "
            "Horizontal dashed lines mark the conventional therapeutic AUC range boundaries of "
            "30 and 60 mg.h/L. The plots show that low body weight and higher total exposure shift "
            "patients toward higher toxicity probabilities, whereas the therapeutic-index panel "
            "summarizes the balance between immunosuppressive efficacy and adverse-event risk."
        ),
    },
]


def add_supplementary_figures(doc):
    """Append supplementary figures in fixed numeric order with detailed legends."""
    doc.add_page_break()
    doc.add_heading('Supplementary Figures', level=1)

    for idx, fig in enumerate(SUPPLEMENTARY_FIGURES):
        if idx > 0:
            doc.add_page_break()

        doc.add_heading(f"Figure {fig['number']}. {fig['title']}", level=2)
        fig_path = os.path.join(FIG_DIR, fig["filename"])
        if os.path.exists(fig_path):
            picture = doc.add_picture(fig_path, width=Inches(6.5))
            picture.alignment = WD_ALIGN_PARAGRAPH.CENTER

        add_figure_legend(doc, fig["legend"])


# Main-figure legends. The four main figures are supplied to the journal as
# separate files (not embedded inline); their legends are collected into a
# dedicated "Figure Legends" section near the end of the manuscript.
MAIN_FIGURE_LEGENDS = [
    (
        'Figure 1. Population pharmacokinetics under matched contemporary-practice assumptions. '
        'The figure summarizes the semi-mechanistic PBPK framework and simulated steady-state '
        'MPA exposure in Western and Indian virtual transplant populations receiving MMF 1000 mg '
        'twice daily. Panels show the model structure, total MPA AUC_0-12h distributions, free '
        'MPA AUC_0-12h distributions, therapeutic target attainment using the 30-60 mg.h/L total '
        'AUC window, the relationship between body weight and AUC, and interindividual variability. '
        'Both cohorts were simulated with matched albumin and predominantly tacrolimus co-therapy; '
        'therefore, the population exposure difference reflects lower body weight and the exploratory '
        'UGT activity assumptions. Indian patients showed higher total and free MPA exposure and a '
        'larger fraction above the therapeutic AUC window at the same fixed dose.'
    ),
    (
        'Figure 2. Dose optimization for Indian transplant recipients. '
        'The figure compares fixed and weight-based MMF dosing strategies in the Indian virtual '
        'population, including standard 1000 mg twice daily, reduced fixed dosing, and 12 mg/kg '
        'twice daily rounded to available tablet strengths. Panels show therapeutic target '
        'attainment, total MPA AUC distributions, the attenuation of the body weight-AUC relationship '
        'with weight-based dosing, the proposed practical dose nomogram, and summary PK/PD metrics. '
        'The 12 mg/kg strategy reduced systematic overexposure in lower-weight patients, increased '
        'the proportion within the 30-60 mg.h/L AUC target window, and preserved a clinically '
        'implementable tablet-based dosing scheme.'
    ),
    (
        'Figure 3. Linked PK/PD outcomes and efficacy-toxicity tradeoff. '
        'Free MPA concentrations from the PBPK simulations were linked to IMPDH-II inhibition '
        'using a sigmoidal E_max model and then propagated to simulated clinical outcome probabilities. '
        'Panels show the IMPDH concentration-response relationship, population distributions of '
        'average IMPDH inhibition, exposure-zone classification, rejection risk, gastrointestinal '
        'toxicity risk, leukopenia risk, the efficacy-toxicity landscape, and the composite therapeutic '
        'index. Standard 1000 mg twice-daily dosing in Indian patients increased IMPDH inhibition and '
        'slightly reduced simulated rejection risk, but this came at the cost of higher toxicity. '
        'Weight-based dosing shifted the population toward a more favorable balance by lowering '
        'toxicity while maintaining adequate pharmacodynamic effect.'
    ),
    (
        'Figure 4. Model validation against published clinical PK studies. '
        'Predicted steady-state total MPA AUC values were compared with observed AUC values from '
        'five published renal-transplant PK datasets spanning Western, Chinese, Thai, tacrolimus, '
        'and cyclosporine co-therapy settings. Panels show predicted versus observed AUC, fold error '
        'by study, the validation performance summary, and predicted AUC distributions overlaid with '
        'reported observed means. The overall geometric mean fold error was 1.23; three of five '
        'studies met the stringent 0.8-1.25 fold criterion and all studies were within two-fold. '
        'Dedicated Indian renal-transplant AUC cohorts were discussed separately and reserved for '
        'future external validation because dosing, sampling time, and post-transplant timing require '
        'harmonization before direct GMFE comparison.'
    ),
]


def add_main_figure_legends(doc):
    """Emit the main-figure legends in a dedicated 'Figure Legends' section."""
    doc.add_heading('Figure Legends', level=1)
    for text in MAIN_FIGURE_LEGENDS:
        add_figure_legend(doc, text)


# ================================================================
# READ SIMULATION RESULTS FROM GENERATED TABLES
# ================================================================

def read_results():
    """Read pre-generated table files for data to embed in manuscript."""
    results = {}
    t2_path = os.path.join(FIG_DIR, 'Table2_PKPD_Results.txt')
    if os.path.exists(t2_path):
        with open(t2_path) as f:
            results['table2'] = f.read()
    t1_path = os.path.join(FIG_DIR, 'Table1_Demographics.txt')
    if os.path.exists(t1_path):
        with open(t1_path) as f:
            results['table1'] = f.read()
    return results


# ================================================================
# MANUSCRIPT BUILDER
# ================================================================

def create_manuscript():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0
    style.paragraph_format.space_after = Pt(0)

    # ================================================================
    # TITLE PAGE
    # ================================================================
    for _ in range(6):
        doc.add_paragraph('')

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        'Quantitative Systems Pharmacology Analysis of Mycophenolic Acid Exposure '
        'in Indian Transplant Recipients: Role of Body Weight and Pharmacogenomic Variability'
    )
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'

    doc.add_paragraph('')
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Running Title: Weight-Driven MPA Overexposure in Indian Patients')
    run.italic = True
    run.font.size = Pt(11)

    doc.add_paragraph('')

    # ---- Author block (CRediT-compliant authorship) ----
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (name, sup) in enumerate([
        ('Smita Pattanaik', '1'),
        ('Deepesh B Kenwar', '2'),
        ('Sarbpreet Singh', '2'),
        ('Ashish Sharma', '2'),
        ('Abhishek Anil', '3'),
        ('Anand Srinivasan', '3,*'),
    ]):
        authors.add_run(('' if i == 0 else ', ') + name)
        authors.add_run(sup).font.superscript = True

    for num, aff in [
        ('1', 'Department of Pharmacology, Postgraduate Institute of Medical Education and Research, '
              'Chandigarh, India'),
        ('2', 'Department of Renal Transplant Surgery, Postgraduate Institute of Medical Education and '
              'Research, Chandigarh, India'),
        ('3', 'Department of Pharmacology, All India Institute of Medical Sciences, Bhubaneswar, India'),
    ]:
        ap = doc.add_paragraph()
        ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ap.add_run(num).font.superscript = True
        ap.add_run(' ' + aff)

    corr = doc.add_paragraph()
    corr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    corr.add_run('*Corresponding author: ').bold = True
    corr.add_run('Anand Srinivasan, Department of Pharmacology, All India Institute of Medical Sciences, '
                 'Bhubaneswar, India. Email: anandsrinivasan@aiimsbhubaneswar.edu.in')

    emails = doc.add_paragraph()
    emails.alignment = WD_ALIGN_PARAGRAPH.CENTER
    emails.add_run(
        'Author email addresses: Smita Pattanaik (pattanaik.smita2018@gmail.com); '
        'Deepesh B Kenwar (deepesh.doc@gmail.com); Sarbpreet Singh (drsarbpreet@yahoo.com); '
        'Ashish Sharma (ashishpgi@gmail.com); Abhishek Anil (drabhishekanil@gmail.com); '
        'Anand Srinivasan (anandsrinivasan@aiimsbhubaneswar.edu.in).'
    )

    doc.add_page_break()

    # ================================================================
    # ABSTRACT
    # ================================================================
    h = doc.add_heading('Abstract', level=1)
    h.runs[0].font.size = Pt(14)

    abstract_sections = {
        'Background': (
            'Mycophenolate mofetil (MMF) dosing at 1 g twice daily was established in Western populations '
            'with a mean body weight of ~78 kg. Indian renal transplant recipients are often leaner, and recent '
            'Indian AUC studies report frequent MPA overexposure with fixed MMF dosing. The extent to which body '
            'weight alone can drive this overexposure has not been quantitatively isolated.'
        ),
        'Methods': (
            'We developed a semi-mechanistic PBPK/PD model for mycophenolic acid (MPA) incorporating '
            'two-compartment disposition, UGT-mediated metabolism, enterohepatic recirculation, OATP1B1/1B3 '
            'hepatic uptake, and allometric scaling. Virtual populations (n = 500 each) were generated '
            'under a matched contemporary-practice scenario, with serum albumin (4.0 \u00b1 0.4 g/dL) and predominantly '
            'tacrolimus co-therapy (~93%) held similar in both Western and Indian cohorts. The pharmacodynamic '
            'layer linked free MPA to IMPDH-II inhibition (IC\u2085\u2080 = 0.15 mg/L, \u03b3 = 1.5) and '
            'clinical outcome probabilities. The model was validated against five published clinical studies '
            '(GMFE = 1.23).'
        ),
        'Results': (
            'Under contemporary clinical conditions, Indian patients at standard 1 g BID showed 35% higher '
            'total MPA AUC (75.1 vs 55.8 mg\u00b7h/L) and 36% higher free MPA AUC (2.40 vs 1.76 mg\u00b7h/L) '
            'compared to Western counterparts, with greater '
            'interindividual variability (CV% 36.2% vs 30.3%). Factor decomposition confirmed that body weight '
            'and UGT enzyme activity accounted for the entire 19.3 mg\u00b7h/L AUC difference between populations, '
            'while albumin and CNI type contributed negligibly. Variance decomposition '
            'identified UGT1A9 activity as the dominant source of interindividual variability (73% of '
            'explained variance), with the higher Indian CV% driven by wider pharmacogenomic diversity '
            'in UGT enzyme activity (SD 0.18 vs 0.15) and amplified by the nonlinear weight\u2013AUC '
            'relationship (1.68\u00d7 steeper at Indian mean weight). Overexposure (AUC > 60 mg\u00b7h/L) affected '
            '68.0% of Indian vs 34.4% of Western patients. IMPDH inhibition was correspondingly higher '
            '(55.5% vs 45.7%), driving increased GI toxicity (50.5% vs 33.3%) and leukopenia (35.3% vs 21.5%), '
            'with a lower composite '
            'therapeutic index (0.156 vs 0.215). Weight-based dosing at 12 mg/kg BID restored target attainment '
            'to 66.8%, reduced overexposure to 27.2%, '
            'and improved the therapeutic index to 0.219, exceeding Western standard dosing.'
        ),
        'Conclusions': (
            'Under a matched albumin/CNI contemporary-practice scenario, body weight emerges as the principal '
            'modifiable determinant of MPA overexposure in Indian transplant patients. Weight-based dosing '
            'at 12 mg/kg BID provides a mechanistically justified, clinically implementable '
            'strategy to optimize the efficacy\u2013toxicity balance in this population.'
        ),
    }

    for section_name, text in abstract_sections.items():
        p = doc.add_paragraph()
        run_label = p.add_run(f'{section_name}: ')
        run_label.bold = True
        run_label.font.size = Pt(12)
        run_text = p.add_run(text)
        run_text.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(
        'Keywords: mycophenolic acid, body weight, pharmacokinetics, PBPK modeling, Indian population, '
        'dose optimization, transplantation, IMPDH inhibition, allometric scaling, therapeutic drug monitoring'
    )
    run.italic = True
    run.font.size = Pt(11)

    doc.add_page_break()

    # ================================================================
    # INTRODUCTION
    # ================================================================
    doc.add_heading('1. Introduction', level=1)

    doc.add_paragraph(
        'Mycophenolate mofetil (MMF) is a cornerstone immunosuppressant in solid organ transplantation [1]. '
        'As a prodrug of mycophenolic acid (MPA), it selectively inhibits inosine monophosphate '
        'dehydrogenase type II (IMPDH-II), the rate-limiting enzyme in de novo purine synthesis '
        'preferentially used by activated T and B lymphocytes [2].'
    )

    p = doc.add_paragraph('The therapeutic window for MPA is well established [1,3]: ')
    add_inline_math(p, eq_auc_target)
    run = p.add_run(
        ' (total MPA) balances efficacy against acute rejection with acceptable toxicity. '
        'However, the standard fixed dose of 1 g BID was established in clinical trials conducted '
        'predominantly in Western populations with a mean body weight of 75\u201380 kg [1,4].'
    )

    doc.add_paragraph(
        'Indian renal transplant cohorts provide direct evidence that fixed-dose MMF can produce high '
        'MPA exposure: a recent 120-patient Indian study reported mean MPA AUC\u2080\u208b\u2081\u2082\u2095 '
        'of 63.7 \u00b1 23.1 mg\u00b7h/L, with 55.8% of patients above 60 mg\u00b7h/L, and an earlier '
        'Indian pilot study also supported an AUC target range of 30\u201360 mg\u00b7h/L [5,6]. '
        'The recent cohort also emphasized that Indian renal allograft recipients were lean '
        '(BMI 20.8 \u00b1 3.8 kg/m\u00b2), making fixed adult MMF dosing a plausible contributor to '
        'overexposure [5]. Earlier literature identified serum albumin, calcineurin inhibitor '
        '(CNI) co-therapy, enterohepatic recirculation, and pharmacogenomic variation as potential '
        'modifiers of MPA exposure [7,8,9,10,11]. In the present analysis, albumin and CNI '
        'co-therapy were intentionally held comparable across virtual cohorts as a model-design '
        'choice. This design allowed the residual contribution of body weight and UGT activity '
        'assumptions to MPA overexposure to be quantified directly.'
    )

    doc.add_paragraph(
        'We therefore developed a PBPK/PD simulation framework in which both Western and Indian '
        'virtual populations reflect a matched contemporary-practice scenario\u2014comparable serum albumin and '
        'predominantly tacrolimus-based co-therapy\u2014with body weight as the primary distinguishing '
        'physiological parameter. This clinically representative design enables quantification of '
        'weight-driven MPA overexposure and evaluation of whether weight-based dosing can resolve it.'
    )

    doc.add_paragraph(
        'The objectives of this study were to: (1) quantify the extent of MPA overexposure in Indian '
        'patients under contemporary clinical conditions, where body weight is the primary differentiating '
        'factor; (2) characterize the increased interindividual variability in Indian patients; (3) evaluate the '
        'pharmacodynamic consequences (IMPDH inhibition, toxicity, therapeutic index); and (4) validate '
        'weight-based dosing at 12 mg/kg BID as an optimal strategy.'
    )

    # ================================================================
    # METHODS
    # ================================================================
    doc.add_heading('2. Methods', level=1)

    doc.add_heading('2.1 Pharmacokinetic Model', level=2)
    doc.add_paragraph(
        'A semi-mechanistic PBPK model was constructed for MPA incorporating: (1) first-order absorption '
        'with lag time (k\u2090 = 4.0 /h, t\u2097\u2090\u1d4d = 0.25 h), (2) two-compartment distribution '
        '(V\u1d9c = 50 L, V\u209a = 150 L, Q = 30 L/h), (3) hepatic metabolism dominated by UGT1A9, '
        'forming mycophenolate glucuronide, with UGT2B7 contribution to acyl-glucuronide formation [12], '
        'modeled using the well-stirred equation, '
        '(4) enterohepatic recirculation (EHC) via '
        'ABCC2/MRP2-mediated biliary excretion [8], and (5) renal elimination of MPAG proportional to GFR.'
    )

    p_eq = doc.add_paragraph('The well-stirred hepatic clearance model:')
    create_display_equation(doc, eq_well_stirred, number='1')

    doc.add_paragraph(
        'where CLₕ is hepatic clearance (L/h), Qₕ is hepatic blood flow (L/h), fᵤ is the unbound '
        'fraction of MPA in plasma, and CLᵢₙₜ is the total intrinsic clearance (L/h).'
    )

    p = doc.add_paragraph('The hepatic extraction ratio:')
    create_display_equation(doc, eq_extraction_ratio, number='2')

    p = doc.add_paragraph(
        'Because MPA has a low extraction ratio (E \u2248 0.13), hepatic clearance is highly sensitive '
        'to changes in protein binding [7]:'
    )
    create_display_equation(doc, eq_clh_approx, number='3')

    doc.add_paragraph(
        'The free fraction was modeled dynamically as a function of albumin and circulating MPAG '
        '(competitive displacement) [7]:'
    )
    create_display_equation(doc, eq_free_fraction, number='4')

    p = doc.add_paragraph()
    run = p.add_run(
        'where f\u1d64,ref = 0.03 at albumin 4.0 g/dL. Both populations were modeled with albumin '
        '4.0 \u00b1 0.4 g/dL as a matched scenario assumption rather than as a population-level '
        'difference between Western and Indian recipients.'
    )

    doc.add_paragraph(
        'Allometric scaling related PK parameters from the reference 70 kg adult to individual patients [13,14]:'
    )
    create_display_equation(doc, eq_allometric_volume, number='5a')
    create_display_equation(doc, eq_allometric_clearance, number='5b')

    doc.add_paragraph(
        'Liver weight was scaled with exponent 0.86 [13]. The cyclosporine\u2013MPA interaction was modeled '
        'through dual scenario mechanisms informed by published EHC and co-medication modeling: '
        '(1) reduced biliary MPAG excretion/EHC and (2) reduced hepatic MPAG uptake [8]. The '
        '40% EHC reduction and 60% uptake reduction were model assumptions used to represent a '
        'cyclosporine scenario, not directly estimated population effects. In this study, both '
        'populations were modeled with ~93% tacrolimus use as a matched contemporary-practice '
        'scenario; this assumption is consistent with the tacrolimus-based regimen used in a recent '
        'Indian renal-transplant exposure cohort [5].'
    )

    doc.add_paragraph('The central compartment mass balance:')
    create_display_equation(doc, eq_ode_central, number='6')

    doc.add_paragraph(
        'The system was solved using an adaptive-step LSODA integrator (rtol = 10\u207b\u2078, '
        'atol = 10\u207b\u00b9\u2070) over 14 BID doses to achieve steady state.'
    )

    # 2.2 PD Model
    doc.add_heading('2.2 Pharmacodynamic Model', level=2)
    doc.add_paragraph(
        'IMPDH type II inhibition was modeled using a sigmoidal E\u2098\u2090\u2093 equation:'
    )
    create_display_equation(doc, eq_emax_impdh, number='7')

    doc.add_paragraph(
        'with IC\u2085\u2080 = 0.15 mg/L free MPA and Hill coefficient \u03b3 = 1.5, calibrated against '
        'published IMPDH activity and variability data [15,16]. These PD values are phenomenological '
        'calibration parameters rather than direct ethnic or population-specific estimates. '
        'Identical PD parameters were applied to both populations, as there is no evidence '
        'for ethnic differences in IMPDH-II sensitivity to MPA.'
    )

    doc.add_paragraph(
        'Clinical outcome probabilities were estimated using phenomenological exposure\u2013response '
        'relationships informed by MPA TDM and Indian renal-transplant exposure studies [3,5,6,17]. '
        'These included rejection probability (inversely related to IMPDH inhibition and AUC), GI toxicity (E\u2098\u2090\u2093 '
        'model with AUC\u2085\u2080 = 70 mg\u00b7h/L), leukopenia (related to IMPDH inhibition with '
        'IMPDH\u2085\u2080 = 0.65), and infection risk (lymphocyte nadir-dependent). '
        'A composite therapeutic index was defined:'
    )
    create_display_equation(doc, eq_therapeutic_index, number='8')
    doc.add_paragraph('(Pₐₙᵧ is the composite adverse event probability)')

    doc.add_paragraph(
        'Interindividual variability was quantified using the coefficient of variation:'
    )
    create_display_equation(doc, eq_cv, number='9')

    # 2.3 Population Design
    doc.add_heading('2.3 Virtual Population Design', level=2)
    doc.add_paragraph(
        'Virtual populations of 500 patients each were generated to reflect contemporary '
        'transplant demographics under a matched albumin/CNI scenario. Both populations shared '
        'the following model-design assumptions: serum albumin 4.0 \u00b1 0.4 g/dL (truncated '
        'normal, range 2.5\u20135.5 g/dL) and tacrolimus co-therapy at ~93%. These shared '
        'parameters were used to isolate body weight and UGT activity rather than to claim that '
        'albumin or CNI use is identical across all real-world cohorts. '
        'Differing parameters were Western body weight 78 \u00b1 15 kg and Indian body weight '
        '58 \u00b1 12 kg; Western eGFR 55 \u00b1 18 mL/min, Indian eGFR 50 \u00b1 18 mL/min; '
        'Indian UGT1A9 activity 0.95 \u00b1 0.18 (vs 1.0 \u00b1 0.15), Indian UGT2B7 activity '
        '0.92 \u00b1 0.20 (vs 1.0 \u00b1 0.15). '
        'The Indian weight distribution was chosen to represent a lower-weight/lean recipient '
        'scenario consistent with Indian and Asian transplant exposure reports [4,5,6]. '
        'UGT activities were sampled from log-normal distributions reflecting pharmacogenomic '
        'variability. The modest mean reductions modeled in Indian populations reflect an exploratory '
        'prior rather than an established effect-size estimate, as the direct genotype\u2013activity data '
        'specific to South Asian MPA metabolism are limited, and reported UGT2B7 c.802C>T (rs7439366) '
        'T-allele frequencies in South Indian cohorts are approximately 37–40% [18, 19], below the '
        '≈49–52% reported in European cohorts. The robustness of conclusions to this assumption '
        'is examined in the parameter-sensitivity analysis (Figure S1).'
    )

    # 2.4 Dosing
    doc.add_heading('2.4 Dosing Strategies', level=2)
    doc.add_paragraph(
        'Four strategies were compared: (1) Western standard 1000 mg BID, (2) Indian standard 1000 mg BID, '
        '(3) Indian reduced fixed-dose 750 mg BID, and (4) Indian weight-based 12 mg/kg BID rounded to '
        'the nearest 250 mg tablet strength. The weight-based strategy was derived from prior dose '
        'optimization analysis targeting a midpoint AUC of 45 mg\u00b7h/L.'
    )

    # 2.5 Validation
    doc.add_heading('2.5 Model Validation', level=2)
    doc.add_paragraph(
        'The model was validated against five published clinical PK studies spanning Western, Chinese, '
        'and Thai populations with both cyclosporine and tacrolimus co-therapy [7\u201311]. Predictive performance '
        'was assessed using:'
    )
    create_display_equation(doc, eq_gmfe, number='10')
    doc.add_paragraph(
        'Acceptance criteria: GMFE < 2.0 overall, individual predictions ideally within 0.8\u20131.25 fold.'
    )

    # 2.6 Sensitivity Analysis
    doc.add_heading('2.6 Sensitivity Analysis: Factor Decomposition', level=2)
    doc.add_paragraph(
        'To quantify the relative contribution of each physiological parameter to the '
        'population difference in MPA exposure, a factor decomposition analysis was performed. '
        'Because albumin and CNI co-therapy were deliberately matched between the Western and '
        'Indian virtual cohorts, the primary decomposition focused on the remaining differentiating '
        'parameters: body weight, eGFR, and UGT/ABCC2 activity assumptions. For each scenario, the '
        'percentage change in mean AUC from the Indian baseline was calculated. A companion sensitivity '
        'check confirmed that the matched albumin and CNI assumptions made negligible contribution to '
        'the population AUC difference (Figure S2).'
    )

    # ================================================================
    # RESULTS
    # ================================================================
    doc.add_heading('3. Results', level=1)

    doc.add_heading('3.1 Population Demographics', level=2)
    doc.add_paragraph(
        'The generated virtual populations reflected contemporary clinical demographics '
        '(Table 1). Both populations had mean albumin of 4.0 g/dL and ~93% tacrolimus use by design. '
        'The primary physiological difference between populations was body weight: '
        '78.1 \u00b1 14.1 kg (Western) vs 59.1 \u00b1 10.8 kg '
        '(Indian), resulting in a 35% higher effective mg/kg dose in Indian patients at the same '
        'fixed 1 g dose.'
    )

    # Table 1
    doc.add_paragraph('')
    t1_title = doc.add_paragraph()
    run = t1_title.add_run('Table 1. Virtual Population Demographics (Mean \u00b1 SD)')
    run.bold = True
    run.font.size = Pt(10)

    table1 = doc.add_table(rows=1, cols=3)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.style = 'Light Grid Accent 1'
    for i, text in enumerate(['Parameter', 'Western (n = 500)', 'Indian (n = 500)']):
        cell = table1.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    t1_data = [
        ('Body weight (kg)', '78.1 \u00b1 14.1', '59.1 \u00b1 10.8'),
        ('Serum albumin (g/dL)', '4.0 \u00b1 0.4', '4.0 \u00b1 0.4'),
        ('eGFR (mL/min)', '55.0 \u00b1 16.7', '50.8 \u00b1 16.6'),
        ('Tacrolimus use (%)', '93', '93'),
        ('UGT1A9 activity', '1.01 \u00b1 0.15', '0.96 \u00b1 0.18'),
        ('UGT2B7 activity', '1.01 \u00b1 0.15', '0.94 \u00b1 0.18'),
        ('Effective dose (mg/kg)', '13.2 \u00b1 2.3', '17.8 \u00b1 3.4'),
    ]
    for row_data in t1_data:
        add_table_row(table1, row_data)

    doc.add_paragraph('')

    # 3.2 PK Comparison
    doc.add_heading('3.2 Pharmacokinetic Comparison', level=2)
    doc.add_paragraph(
        'Indian patients demonstrated significantly higher MPA exposure at standard 1 g BID '
        'dosing (Figure 1, Table 2). With albumin and tacrolimus co-therapy matched by design, '
        'the AUC difference was attributable to lower body weight and modestly '
        'reduced UGT activity.'
    )

    doc.add_paragraph(
        'Indian patients exhibited substantially greater interindividual variability in AUC '
        '(CV% 36.2% vs 30.3%). A formal variability decomposition analysis (Figure S3) identified '
        'UGT1A9 enzyme activity as the dominant source of AUC variability in both populations, '
        'accounting for 73% of the explained variance in Indian patients (vs 68% in Western). '
        'Body weight contributed 19% in both populations. However, partial correlation analysis '
        'revealed that when other covariates are controlled, body weight has a strong independent '
        'effect on AUC (partial r = \u22120.75 Indian, \u22120.82 Western), comparable to UGT1A9 '
        '(partial r = \u22120.92 Indian, \u22120.93 Western). The higher Indian CV% is attributable '
        'to two reinforcing mechanisms: (1) wider UGT1A9 activity distribution (SD 0.18 vs 0.15 on '
        'log-scale, modeled as an exploratory pharmacogenomic-diversity prior informed indirectly by '
        'South Asian UGT polymorphism data [18]), and '
        '(2) the nonlinear (hyperbolic) weight\u2013AUC relationship, whereby AUC \u221d BW\u207b\u2070\u00b7\u2077\u2075 [13,14], '
        'making the sensitivity of AUC to weight 1.68\u00d7 steeper at the Indian mean of 58 kg '
        'compared to the Western mean of 78 kg (Figure S3E\u2013F). Thus, identical absolute '
        'variability in body weight translates into greater AUC variability at lower body weights.'
    )

    # Figure 1 supplied as a separate file; legend relocated to the Figure Legends
    # section at the end of the document (see add_main_figure_legends).

    # 3.3 Sensitivity Analysis
    doc.add_heading('3.3 Sensitivity Analysis: Factor Decomposition', level=2)
    doc.add_paragraph(
        'Factor decomposition revealed that body weight and UGT enzyme activity accounted for '
        'the entirety of the AUC difference between populations (Figure S2). Replacing Indian '
        'body weight and UGT parameters with Western values reduced the mean total AUC by '
        '19.3 mg\u00b7h/L, fully bridging the gap from the Indian level (75.1 mg\u00b7h/L) to the '
        'Western level (55.8 mg\u00b7h/L). In contrast, albumin and CNI type contributed a net '
        'change of approximately 0 mg\u00b7h/L to total AUC, confirming their negligible role '
        'within the matched scenario where these parameters were held comparable between populations.'
    )

    doc.add_paragraph(
        'For free MPA AUC, a similar pattern was observed: body weight and UGT differences '
        'accounted for +0.641 mg\u00b7h/L of the total +0.64 mg\u00b7h/L difference (Indian '
        '2.40 vs Western 1.76 mg\u00b7h/L). Within the body weight and UGT contribution, '
        'body weight was the dominant factor: swapping body weight alone to Western values '
        'reduced total AUC by approximately 20%, whereas UGT activity differences contributed '
        'a modest additional reduction. These findings confirm that body weight is the primary '
        'modifiable determinant of the population-level AUC difference.'
    )

    # 3.4 Dose Optimization
    doc.add_heading('3.4 Dose Optimization', level=2)
    doc.add_paragraph(
        'Weight-based dosing at 12 mg/kg BID was the most effective strategy for Indian patients '
        '(Figure 2). This approach effectively decoupled AUC from body weight, eliminating the systematic '
        'overexposure in lower-weight patients. The derived nomogram provides clinically implementable '
        'recommendations (contiguous weight bins), that are 500 mg BID for patients < 50 kg, 750 mg BID for 50\u2013<75 kg, and 1000 mg BID '
        'for \u2265 75 kg.'
    )

    # Figure 2 supplied as a separate file; legend relocated to the Figure Legends
    # section at the end of the document (see add_main_figure_legends).

    # Table 2
    t2_title = doc.add_paragraph()
    run = t2_title.add_run('Table 2. Key Pharmacokinetic and Pharmacodynamic Outcomes')
    run.bold = True
    run.font.size = Pt(10)

    table2 = doc.add_table(rows=1, cols=5)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.style = 'Light Grid Accent 1'
    t2_headers = ['Metric', 'Western\n1000 mg', 'Indian\n1000 mg', 'Indian\n750 mg', 'Indian\n12 mg/kg']
    for i, text in enumerate(t2_headers):
        cell = table2.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    t2_rows = [
        ('Total AUC (mg.h/L)', '55.8 +/- 16.9', '75.1 +/- 27.2', '57.0 +/- 20.7', '52.6 +/- 17.8'),
        ('Free AUC (mg.h/L)', '1.76 +/- 0.51', '2.40 +/- 0.86', '1.80 +/- 0.64', '1.65 +/- 0.55'),
        ('Avg IMPDH inhibition', '0.457 +/- 0.100', '0.555 +/- 0.120', '0.459 +/- 0.119', '0.433 +/- 0.111'),
        ('% Target (30-60 mg.h/L)', '62.2%', '31.0%', '56.8%', '66.8%'),
        ('% Overexposed (>60 mg.h/L)', '34.4%', '68.0%', '38.4%', '27.2%'),
        ('CV% Total AUC', '30.3%', '36.2%', '36.4%', '33.8%'),
        ('P(rejection)', '16.4%', '13.4%', '16.8%', '17.7%'),
        ('P(GI toxicity)', '33.3%', '50.5%', '34.3%', '29.8%'),
        ('P(leukopenia)', '21.5%', '35.3%', '22.4%', '19.0%'),
        ('P(infection)', '6.7%', '11.6%', '7.5%', '6.6%'),
        ('P(any adverse)', '48.7%', '67.2%', '49.3%', '44.2%'),
        ('Therapeutic index', '0.215 +/- 0.054', '0.156 +/- 0.078', '0.206 +/- 0.062', '0.219 +/- 0.053'),
        ('PD zone: Therapeutic', '93.2%', '87.2%', '88.2%', '87.6%'),
        ('PD zone: Overexposed', '0.8%', '11.4%', '2.4%', '0.8%'),
        ('PD zone: Underexposed', '6.0%', '1.4%', '9.4%', '11.6%'),
    ]
    for row in t2_rows:
        cells = table2.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run(text)
            run.font.size = Pt(7)
            run.font.name = 'Times New Roman'
            cells[i].paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )

    t2_note = doc.add_paragraph()
    run = t2_note.add_run('Values are mean +/- SD unless otherwise indicated.')
    run.font.size = Pt(9)
    run.italic = True

    t2_note = doc.add_paragraph()
    run = t2_note.add_run(
        'Abbreviations: AUC, area under the concentration-time curve; IMPDH, inosine '
        'monophosphate dehydrogenase; MMF, mycophenolate mofetil; PD, pharmacodynamic; '
        'CV, coefficient of variation. The first overexposure row is based on total MPA '
        'AUC >60 mg.h/L; PD-zone rows use the clinical-zone classification from the simulation output.'
    )
    run.font.size = Pt(8)
    run.italic = True

    doc.add_paragraph('')

    # 3.4 PD Outcomes
    doc.add_heading('3.5 Pharmacodynamic Outcomes', level=2)
    doc.add_paragraph(
        'The linked PD analysis revealed that the PK overexposure in Indian patients translates directly '
        'into clinically meaningful differences in outcome probabilities (Figure 3). Higher free MPA '
        'concentrations drove greater IMPDH-II inhibition, while reducing rejection risk, '
        'substantially increased toxicity. The composite therapeutic index was lower in Indian patients '
        'at standard dosing, reflecting an unfavorable efficacy\u2013toxicity balance. '
        'Representative steady-state concentration\u2013time and IMPDH-II inhibition profiles across the '
        'four dosing strategies are shown in Figure S4.'
    )

    doc.add_paragraph(
        'Weight-based dosing at 12 mg/kg BID optimally rebalanced the efficacy\u2013toxicity tradeoff. '
        'The therapeutic index improved to levels exceeding Western standard dosing, primarily through '
        'substantial reductions in GI toxicity and leukopenia. The efficacy\u2013toxicity landscape '
        '(Figure 3G) visually demonstrates the shift toward the ideal zone. The dependence of GI '
        'toxicity, leukopenia, and therapeutic index on body weight and total MPA exposure is '
        'illustrated in Figure S5.'
    )

    # Figure 3 supplied as a separate file; legend relocated to the Figure Legends
    # section at the end of the document (see add_main_figure_legends).

    # 3.5 Validation
    doc.add_heading('3.6 Model Validation', level=2)
    doc.add_paragraph(
        'The model demonstrated acceptable predictive performance (Figure 4, Table 3). The GMFE was 1.23, '
        'well within the regulatory acceptance threshold of 2.0. Three of five validation studies fell within '
        'the stringent 0.8\u20131.25 fold criterion, and all five were within 2-fold. Validation against '
        'Shaw et al. [20] (Western/tacrolimus, FE = 1.20), van Hest et al. [10] (Western/CsA, FE = 1.37), '
        'Zicheng et al. [21] (Chinese/CsA, FE = 1.02), Pithukpakorn et al. [9] (Thai/CsA, FE = 0.76), and '
        'Yau et al. [4] (Asian/tacrolimus, FE = 1.09) confirmed acceptable accuracy across diverse '
        'populations and CNI regimens.'
    )

    doc.add_paragraph(
        'Dedicated Indian renal-transplant AUC cohorts are now available but were not included in '
        'the prespecified GMFE validation set. Singh et al. reported AUC\u2080\u208b\u2081\u2082\u2095 at '
        '4 weeks in 120 Indian live-related renal transplant recipients receiving tacrolimus-based '
        'triple therapy, while Sarangi et al. reported a smaller Indian renal-transplant pilot AUC '
        'study [5,6]. The Koloskoff et al. Indian lupus-nephritis popPK dataset remains relevant '
        'for Indian MPA TDM method development but is not a transplant validation set because the '
        'indication and treatment context differ [22]. External validation against Indian transplant '
        'AUC data after harmonizing dosing, sampling, and post-transplant timing remains a priority.'
    )

    # Table 3
    t3_title = doc.add_paragraph()
    run = t3_title.add_run('Table 3. Model Validation Against Published Clinical Pharmacokinetic Studies')
    run.bold = True
    run.font.size = Pt(10)

    table3 = doc.add_table(rows=1, cols=5)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    table3.style = 'Light Grid Accent 1'
    t3_headers = ['Study', 'Population', 'CNI co-therapy', 'Fold error', 'Within 0.8-1.25 fold']
    for i, text in enumerate(t3_headers):
        cell = table3.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    t3_rows = [
        ('Shaw et al. [20]', 'Western', 'Tacrolimus', '1.20', 'Yes'),
        ('van Hest et al. [10]', 'Western', 'Cyclosporine', '1.37', 'No'),
        ('Zicheng et al. [21]', 'Chinese', 'Cyclosporine', '1.02', 'Yes'),
        ('Pithukpakorn et al. [9]', 'Thai', 'Cyclosporine', '0.76', 'No'),
        ('Yau et al. [4]', 'Asian', 'Tacrolimus', '1.09', 'Yes'),
        ('GMFE (overall)', '—', '—', '1.23', '—'),
    ]
    for row_data in t3_rows:
        add_table_row(table3, row_data)

    t3_note = doc.add_paragraph()
    run = t3_note.add_run(
        'FE, fold error (predicted/observed total MPA AUC₀₋₁₂ₕ); '
        'GMFE, geometric mean fold error; CNI, calcineurin inhibitor. Three of five studies met the '
        '0.8–1.25 fold criterion and all five were within two-fold.'
    )
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'

    # Figure 4 supplied as a separate file; legend relocated to the Figure Legends
    # section at the end of the document (see add_main_figure_legends).

    # ================================================================
    # 3.7 PARAMETER SENSITIVITY ANALYSIS
    # ================================================================
    doc.add_heading('3.7 Parameter Uncertainty Sensitivity Analysis', level=2)

    doc.add_paragraph(
        'A local one-at-a-time sensitivity analysis was performed to assess the robustness of '
        'model predictions to uncertainty in structural parameters. Each of six key PBPK parameters '
        '(f\u1d64,ref: reference unbound fraction of MPA at normal albumin, CLint for UGT1A9: '
        'intrinsic metabolic clearance through UGT1A9, CLint for UGT2B7: intrinsic metabolic '
        'clearance through UGT2B7, central volume V\u1d84, absorption rate k\u2090, and biliary efflux '
        'rate k_bile) was perturbed by '
        '\u00b120% from its nominal value, with simulations performed at the Indian mean-patient '
        'baseline (BW 58 kg, albumin 4.0 g/dL, tacrolimus co-therapy, UGT1A9 0.95, UGT2B7 0.92, '
        'eGFR 50 mL/min; baseline AUC = 73.5 mg\u00b7h/L at 1 g MMF BID steady state).'
    )
    doc.add_paragraph(
        'The results (Figure S1) identify the free fraction at reference albumin (f\u1d64,ref) '
        'as the single most influential parameter (AUC change +21.5%/\u221214.5% at \u00b120% perturbation), '
        'followed by UGT1A9 intrinsic clearance (+15.4%/\u221211.4%). UGT2B7 CL\u1d62\u2099\u209c, '
        'biliary efflux rate, central volume, and absorption rate individually contributed '
        '\u22645% AUC change. Critically, even the largest single-parameter swing (\u223c22%) is '
        'smaller in magnitude than the \u223c35% Indian-vs-Western AUC difference reported at '
        'standard dosing, indicating that the study\u2019s population-level conclusions are '
        'structurally robust to plausible parameter uncertainty. The analysis also confirms that '
        'the modest UGT activity priors assumed for the Indian cohort (\u00b15\u20138% mean shifts) '
        'contribute substantially less to the between-population AUC difference than body weight.'
    )

    # ================================================================
    # DISCUSSION
    # ================================================================
    doc.add_page_break()
    doc.add_heading('4. Discussion', level=1)

    doc.add_paragraph(
        'This study quantifies the contribution of body weight to MPA overexposure in Indian '
        'transplant patients under contemporary clinical conditions. The contemporary-practice '
        'design deliberately holds albumin and CNI co-therapy comparable across populations, thereby '
        'isolating body weight and UGT pharmacogenomic variation as the remaining differentiating '
        'parameters. Within that design, the 35% '
        'higher total AUC (75.1 vs 55.8 mg\u00b7h/L) and 36% higher free AUC '
        '(2.40 vs 1.76 mg\u00b7h/L) in Indian patients at standard 1 g BID dosing '
        'are driven predominantly by the 26% lower mean body weight, with modest additional '
        'contribution from assumed UGT activity differences. This weight-based interpretation is '
        'consistent with Asian weight-adjusted dosing data and recent Indian reports of frequent '
        'fixed-dose MPA overexposure [4,5,6].'
    )

    doc.add_paragraph(
        'The allometric basis of this effect is straightforward [13,14]: at fixed dosing, lower body weight '
        'results in (1) higher mg/kg dose delivery, (2) smaller distribution volumes (V \u221d BW\u00b9\u00b7\u2070), '
        'and (3) reduced clearance (CL \u221d BW\u2070\u00b7\u2077\u2075) that does not fully compensate for '
        'the increased dose-to-weight ratio. For a low extraction ratio drug like MPA (E \u2248 0.13), '
        'hepatic clearance depends on f\u1d64 \u00d7 CL\u1d62\u2099\u209c (Equation 3) [7], where CL\u1d62\u2099\u209c '
        'scales with liver weight (BW\u2070\u00b7\u2078\u2076). The net effect is that '
        'AUC scales approximately inversely with body weight at fixed doses.'
    )

    doc.add_paragraph(
        'A novel finding of this study is the rigorous quantification of the sources of greater '
        'interindividual variability in Indian patients (CV% 36.2% vs 30.3% for total AUC). '
        'The variance decomposition analysis (Figure S3) revealed that UGT1A9 enzyme activity is the '
        'dominant source of AUC variability in both populations (73% of explained variance in Indian, '
        '68% in Western patients), followed by body weight (19% in both). The higher Indian variability '
        'is driven by the wider UGT1A9 activity distribution (SD 0.18 vs 0.15 on log-scale), an '
        'exploratory prior informed indirectly by South Asian UGT polymorphism data (e.g., UGT2B7 '
        'c.802C>T T-allele frequency ≈37–40% in a South Indian non-transplant cohort [18]); direct '
        'South Asian-specific UGT1A9 functional-variant data for MPA remain sparse.'
    )

    doc.add_paragraph(
        'Importantly, the partial correlation analysis disambiguates the confounded roles of UGT1A9 and '
        'body weight. While UGT1A9 dominates overall variance (Pearson r = \u22120.80 with AUC in Indian '
        'patients), body weight has a strong independent effect when other covariates are controlled '
        '(partial r = \u22120.75). This distinction is critical: UGT1A9 variability drives '
        'within-population spread, while body weight drives the between-population mean shift [10].'
    )

    doc.add_paragraph(
        'A second mechanism amplifying Indian AUC variability is the nonlinear weight\u2013exposure '
        'relationship. Because AUC scales as BW\u207b\u2070\u00b7\u2077\u2075 at fixed dosing (a consequence of '
        'clearance scaling as BW\u2070\u00b7\u2077\u2075 while dose remains constant) [13,14], the derivative '
        '|dAUC/dBW| = 0.75 \u00d7 k \u00d7 BW\u207b\u00b9\u00b7\u2077\u2075 is 1.68\u00d7 steeper at '
        'the Indian mean weight (58 kg) compared to the Western mean (78 kg). This means that a given '
        'kilogram of weight variation produces 68% more AUC variation in the Indian weight range. '
        'Since both populations have similar weight CV% (~20%), this mathematical amplification '
        'alone would increase AUC CV% by approximately 3 percentage points in the Indian population, '
        'closely matching the observed difference.'
    )

    doc.add_paragraph(
        'These findings have direct clinical implications. The higher variability in Indian patients '
        'suggests a stronger rationale for routine therapeutic drug monitoring (TDM) and AUC-guided '
        'dose adjustment in this population [3,5,6,17]. Furthermore, the pharmacogenetic architecture '
        'of MPA metabolism differs substantially between South Asian and Western populations in ways '
        'directly relevant to overexposure. The UGT1A9 promoter gain-of-function variants \u2212275T>A '
        'and \u22122152C>T, which increase hepatic UGT1A9 expression, lower MPA exposure by ~20%, '
        'and have been associated with acute rejection in MMF/tacrolimus-treated Western recipients '
        '[23] are essentially absent in South Asians (gnomAD minor-allele frequency \u22481% each, '
        'versus \u22485\u20136% in Europeans and 0% in East Asians [11,19,24]). The reduced-activity '
        'UGT1A9*3 (c.98T>C) allele, which conversely raises MPA exposure, is likewise near-absent in '
        'South Asians (<0.2% vs \u22481.4% in Europeans [19]). Consequently, the common European '
        'variants that shift MPA clearance are not operative at the population level in Indian patients, '
        'who lack the clearance-accelerating genotype that partially protects a subset (~10\u201315%) of '
        'Western recipients from overexposure. By contrast, the intronic UGT1A9 I399C>T variant, which is '
        'reported to influence MPA pharmacokinetics in Japanese renal transplant recipients [24], '
        'reaches its highest global frequency in South Asians (\u224833%, versus \u224830% in Europeans '
        'and \u22483% in East Asians [19]), although its functional impact on MPA exposure remains '
        'unresolved. Together with an intermediate UGT2B7 c.802C>T T-allele frequency (\u224837\u201340% '
        'in South Indian cohorts [18], below the \u224849\u201352% reported in Europeans), these data '
        'indicate that the UGT variants most strongly linked to MPA exposure in Western and East Asian '
        'studies do not transfer directly to South Asians. Pharmacogenomic-guided dosing in Indian '
        'recipients will therefore require South Asian-specific functional-variant discovery rather '
        'than genotyping the established European markers, reinforcing AUC-guided TDM as the pragmatic '
        'near-term strategy [3,17].'
    )

    doc.add_paragraph(
        'The PD consequences of overexposure are clinically significant. Higher IMPDH inhibition (55.5% '
        'vs 45.7%) provides marginally better protection against rejection (13.4% vs 16.4%) but substantially '
        'increases toxicity risk: GI toxicity 50.5% vs 33.3%, leukopenia 35.3% vs 21.5%, and any adverse '
        'event 67.2% vs 48.7%. The therapeutic index (Equation 8) captures this unfavorable balance '
        'quantitatively (0.156 vs 0.215). Weight-based dosing at 12 mg/kg BID restores the therapeutic '
        'index to 0.219 by proportionally reducing the dose for lower-weight patients.'
    )

    doc.add_paragraph(
        'Several limitations should be acknowledged. First, the model was calibrated against Western PK '
        'data and validated against non-Indian transplant PK studies [4,9,10,20,21]; it has not yet been '
        'externally validated against the available Indian renal-transplant AUC cohorts [5,6]. '
        'Koloskoff et al. [22] is an Indian lupus-nephritis popPK study and is therefore relevant '
        'for TDM methodology but not directly interchangeable with transplant validation. Second, the PD '
        'model uses simplified logistic models for clinical outcomes rather than time-to-event analyses. '
        'Third, while UGT pharmacogenomic variability was incorporated as an exploratory prior, '
        'comprehensive South Asian-specific UGT1A9/UGT2B7 functional allele frequency data for MPA '
        'remain limited [11,18]. Fourth, the model uses matched steady-state albumin assumptions and '
        'does not simulate transient early post-transplant albumin changes. '
        'Despite these limitations, the use of clinically '
        'representative population parameters strengthens the translational relevance of the primary '
        'conclusion regarding body weight.'
    )

    # ================================================================
    # CONCLUSIONS
    # ================================================================
    doc.add_heading('5. Conclusions', level=1)
    doc.add_paragraph(
        'This PBPK/PD analysis quantifies the contribution of body weight to MPA overexposure in '
        'Indian transplant patients under contemporary clinical conditions. The contemporary-practice '
        'design isolates body weight and the UGT pharmacogenomic variation as the remaining '
        'differentiating parameters between Western and Indian cohorts. '
        'At standard 1 g BID dosing with clinically '
        'representative albumin and tacrolimus-based regimens, 68% of Indian patients exceed the '
        'therapeutic AUC window, with correspondingly elevated toxicity risk '
        '(67.2% any adverse event vs 48.7%). Weight-based dosing at 12 mg/kg BID (nomogram: 500 mg '
        'for <50 kg, 750 mg for 50\u2013<75 kg, 1000 mg for \u226575 kg) optimizes both target attainment '
        'and the therapeutic index. These findings, robust to \u00b120% parameter perturbation, '
        'provide a mechanistic justification for prospective clinical trials of individualized MPA dosing '
        'in Indian and other low-weight transplant populations.'
    )

    # ================================================================
    # DECLARATIONS / BACK MATTER
    # ================================================================
    doc.add_heading('Author Contributions', level=1)
    credit = doc.add_paragraph()
    _credit_roles = [
        ('Anand Srinivasan', 'Conceptualization, Methodology, Software, Formal analysis, Validation, '
                             'Visualization, Writing – original draft, Supervision, Project administration.'),
        ('Smita Pattanaik', 'Conceptualization, Methodology, Formal analysis, Validation, '
                            'Writing – original draft, Writing – review & editing, Supervision.'),
        ('Deepesh B Kenwar', 'Resources, Investigation, Writing – review & editing.'),
        ('Sarbpreet Singh', 'Resources, Investigation, Writing – review & editing.'),
        ('Ashish Sharma', 'Resources, Investigation, Supervision, Writing – review & editing.'),
        ('Abhishek Anil', 'Data curation, Validation, Visualization, Writing – review & editing.'),
    ]
    for i, (name, roles) in enumerate(_credit_roles):
        credit.add_run(('' if i == 0 else ' ') + name + ': ').bold = True
        credit.add_run(roles)

    doc.add_heading('Funding', level=1)
    doc.add_paragraph(
        'This research did not receive any specific grant from funding agencies in the public, '
        'commercial, or not-for-profit sectors.'
    )

    doc.add_heading('Conflict of Interest', level=1)
    doc.add_paragraph('The authors have no conflicts of interest to declare.')

    doc.add_heading('Data Availability Statement', level=1)
    doc.add_paragraph(
        'The code scripts used to generate the model and simulations reported in this study are '
        'available from the corresponding author upon reasonable request.'
    )

    add_main_figure_legends(doc)

    add_supplementary_figures(doc)

    references = build_references()
    transform_bracket_citations_in_doc(doc, references)

    # Save
    doc.save(OUTPUT_PATH)
    print(f"EndNote-ready manuscript saved: {OUTPUT_PATH}")

    # Emit EndNote-importable RIS sidecar
    write_ris(references, RIS_PATH)
    print(f"RIS reference file saved: {RIS_PATH}")


if __name__ == "__main__":
    create_manuscript()
