# One-off editor: update author block + add CRediT/declarations to the revision docx.
# Edits MPA_QSP_Manuscript_Revision.docx in place (backup: .preauthors.docx).
import os
from docx import Document

HERE = os.path.join(os.path.dirname(__file__), '..', 'outputs', 'manuscript_updated')
PATH = os.path.join(HERE, 'MPA_QSP_Manuscript_Revision.docx')

d = Document(PATH)
paras = d.paragraphs

# ---- locate anchor paragraphs by content (robust to index drift) ----
author_p = runtitle_p = concl_text_p = None
for i, p in enumerate(paras):
    t = p.text.strip()
    if author_p is None and t.startswith('Anand Srinivasan, Smita Pattanaik'):
        author_p = p
    if runtitle_p is None and t.startswith('Running Title'):
        runtitle_p = p
    if t == '5. Conclusions':
        concl_text_p = paras[i + 1]
assert author_p and runtitle_p and concl_text_p, (author_p, runtitle_p, concl_text_p)

AUTHORS = [
    ("Smita Pattanaik", "1"),
    ("Deepesh B Kenwar", "2"),
    ("Sarbpreet Singh", "2"),
    ("Ashish Sharma", "2"),
    ("Abhishek Anil", "3"),
    ("Anand Srinivasan", "3,*"),
]
AFFILS = [
    ("1", "Department of Pharmacology, Postgraduate Institute of Medical Education and Research, Chandigarh, India"),
    ("2", "Department of Renal Transplant Surgery, Postgraduate Institute of Medical Education and Research, Chandigarh, India"),
    ("3", "Department of Pharmacology, All India Institute of Medical Sciences, Bhubaneswar, India"),
]
EMAILS = [
    ("Smita Pattanaik", "pattanaik.smita2018@gmail.com"),
    ("Deepesh B Kenwar", "deepesh.doc@gmail.com"),
    ("Sarbpreet Singh", "drsarbpreet@yahoo.com"),
    ("Ashish Sharma", "ashishpgi@gmail.com"),
    ("Abhishek Anil", "drabhishekanil@gmail.com"),
    ("Anand Srinivasan", "anandsrinivasan@aiimsbhubaneswar.edu.in"),
]

# ---- 1) rewrite the author line in place ----
for r in list(author_p.runs):
    r._element.getparent().remove(r._element)
for idx, (name, sup) in enumerate(AUTHORS):
    author_p.add_run(("" if idx == 0 else ", ") + name)
    author_p.add_run(sup).font.superscript = True

# ---- helper: insert a paragraph (list of (text, fmt) runs) before ref_p ----
def ins_before(ref_p, runs, style=None):
    np = ref_p.insert_paragraph_before(style=style)
    for text, fmt in runs:
        r = np.add_run(text)
        if fmt.get('bold'):   r.bold = True
        if fmt.get('italic'): r.italic = True
        if fmt.get('super'):  r.font.superscript = True
    return np

# ---- 2) affiliations + corresponding author + emails (before running title) ----
ins_before(runtitle_p, [("", {})])
for num, text in AFFILS:
    ins_before(runtitle_p, [(num, {'super': True}), (" " + text, {})])
ins_before(runtitle_p, [("", {})])
ins_before(runtitle_p, [
    ("*Corresponding author: ", {'bold': True}),
    ("Anand Srinivasan, Department of Pharmacology, All India Institute of Medical Sciences, "
     "Bhubaneswar, India. Email: anandsrinivasan@aiimsbhubaneswar.edu.in", {}),
])
ins_before(runtitle_p, [("Author email addresses: "
                         + "; ".join(f"{n} ({e})" for n, e in EMAILS) + ".", {})])
ins_before(runtitle_p, [("", {})])

# ---- 3) back matter after Conclusions body ----
anchor = paras[paras.index(concl_text_p) + 1]  # first spacer after conclusions

def add_section(heading, body_runs):
    h = anchor.insert_paragraph_before(style='Heading 1')
    h.add_run(heading)
    bp = anchor.insert_paragraph_before()
    for text, fmt in body_runs:
        r = bp.add_run(text)
        if fmt.get('bold'):   r.bold = True
        if fmt.get('italic'): r.italic = True
    anchor.insert_paragraph_before()  # trailing spacer

credit_body = [
    ("Anand Srinivasan: ", {'bold': True}),
    ("Conceptualization, Methodology, Software, Formal analysis, Validation, Visualization, "
     "Writing – original draft, Supervision, Project administration. ", {}),
    ("Smita Pattanaik: ", {'bold': True}),
    ("Conceptualization, Methodology, Formal analysis, Validation, Writing – original draft, "
     "Writing – review & editing, Supervision. ", {}),
    ("Deepesh B Kenwar: ", {'bold': True}),
    ("Resources, Investigation, Writing – review & editing. ", {}),
    ("Sarbpreet Singh: ", {'bold': True}),
    ("Resources, Investigation, Writing – review & editing. ", {}),
    ("Ashish Sharma: ", {'bold': True}),
    ("Resources, Investigation, Supervision, Writing – review & editing. ", {}),
    ("Abhishek Anil: ", {'bold': True}),
    ("Data curation, Validation, Visualization, Writing – review & editing.", {}),
]
add_section("Author Contributions", credit_body)
add_section("Funding", [("This research did not receive any specific grant from funding agencies "
                         "in the public, commercial, or not-for-profit sectors.", {})])
add_section("Conflict of Interest", [("The authors have no conflicts of interest to declare.", {})])
add_section("Data Availability Statement", [
    ("The code scripts used to generate the model and simulations reported in this study are "
     "available from the corresponding author upon reasonable request.", {})])

d.save(PATH)
print("saved OK")
